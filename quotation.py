import tkinter as tk
from tkinter import filedialog, messagebox, colorchooser, simpledialog
from PIL import Image, ImageTk
import os
import datetime
import tempfile
import webbrowser
import re
import sys
# import pyodbc
import sqlite3
import shutil 
import json
import urllib.parse
import hashlib
import subprocess
import zipfile
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from ttkbootstrap.widgets import ToastNotification
import ttkbootstrap as tb # Keep tb as an alias for safety

import ui_styles as style
from theme_manager import ThemeManager
# Optional imports with safe handling
docx = None
try:
    import docx
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.shared import OxmlElement
    from docx.table import _Cell
    try:
        from docx.enum.table import WD_ALIGN_VERTICAL
    except ImportError:
        WD_ALIGN_VERTICAL = None
except ImportError:
    pass

reportlab = None
try:
    import reportlab
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image as RLImage
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.utils import ImageReader
    from reportlab.lib.enums import TA_CENTER, TA_RIGHT, TA_LEFT
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
except ImportError:
    pass

qrcode = None
try:
    import qrcode
except ImportError:
    pass

openpyxl = None
try:
    import openpyxl
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
except ImportError:
    pass

class QuotationApp:
    # =========================================================================
    # INITIALIZATION & STARTUP
    # =========================================================================
    def __init__(self, root):
        self.root = root
        self.root.title("Professional Quotation Generator V5.1")
        self.root.geometry("1300x850")
        
        # --- INITIAL STATE ---
        self.current_username = "admin"  # Default before DB init
        self.current_db_id = None
        self.current_dashboard = None
        
        # --- MODERN STYLING ---
        # Initialize theme via centralized ThemeManager
        self.style = ThemeManager.apply_theme(self.root)
        self.style.configure('TEntry', padding=5)
    
        self.init_database()
        
        # --- State Variables ---
        self.header_logo_path = None
        self.middle_logo_path = None
        self.client_logo_path = None
        self.footer_logo_path = None
        self.header_color = "#e0e0e0" 
        self.editing_index = None 
        
        # Settings
        self.currency_var = tk.StringVar(value="PKR")
        self.currency_symbol_var = tk.StringVar(value="Rs.")
        self.gst_rate_var = tk.DoubleVar(value=18.0)
        
        # Logo Sizes
        self.v_logo_size_var = tk.DoubleVar(value=1.5)
        self.m_logo_size_var = tk.DoubleVar(value=1.5)
        self.c_logo_size_var = tk.DoubleVar(value=1.5)
        self.f_logo_size_var = tk.DoubleVar(value=5.0)

        # Logo Alignment Vars
        self.vendor_logo_align_var = tk.StringVar(value="Left")
        self.client_logo_align_var = tk.StringVar(value="Right")
        
        # Text Styling Vars
        self.font_family_var = tk.StringVar(value="Helvetica")
        self.font_size_var = tk.StringVar(value="10")
        
        # Tracking for notifications
        self.last_pdf_pages = 1
        
        # Global Rates for percentage columns (id -> DoubleVar)
        self.global_rates = {}

        # Data
        self.items_data = [] 
        self.extra_fields = [] 
        self.row_colors = {} 
        
        # Auto-Save Feature
        self.auto_save_enabled = tk.BooleanVar(value=True)
        self.last_auto_save_time = None
        self.auto_save_timer = None
        self.current_db_id = None  # Track current quotation ID in database
        self.current_dashboard = None

        # --- CUSTOM COLUMNS SETUP ---
        self.columns_config = [
            {'id': 'sno',   'label': 'S.No', 'width': 35, 'type': 'auto'},
            {'id': 'desc',  'label': 'Description', 'width': 300, 'type': 'text'},
            {'id': 'uom',   'label': 'UOM', 'width': 50, 'type': 'text'},
            {'id': 'qty',   'label': 'Qty', 'width': 60, 'type': 'number'},
            {'id': 'price', 'label': 'Price', 'width': 130, 'type': 'number'},
            {'id': 'amount','label': 'Amount', 'width': 130, 'type': 'calc'},
            {'id': 'gst',   'label': 'Tax', 'width': 90, 'type': 'calc'},
            {'id': 'total', 'label': 'Total', 'width': 130, 'type': 'calc'}
        ]
        self.dynamic_vars = {}

        # --- FINANCIAL TOTALS STATE ---
        self.subtotal_var = tk.StringVar(value="0.00")
        self.tax_total_var = tk.StringVar(value="0.00")
        self.grand_total_var = tk.StringVar(value="0.00")
        
        # Print selection flags
        self.print_subtotal_var = tk.BooleanVar(value=True)
        self.print_tax_var = tk.BooleanVar(value=True)
        self.print_grand_total_var = tk.BooleanVar(value=True)

        # Footer State
        self.footer_align_var = tk.StringVar(value="Center")
        self.footer_text_var = tk.StringVar(value="")
        self.footer_full_width_var = tk.BooleanVar(value=False)
        self.footer_pin_to_bottom_var = tk.BooleanVar(value=True)

        # Header Vars (Standard)
        self.vendor_company_var = tk.StringVar()
        self.vendor_addr_var = tk.StringVar()
        self.vendor_contact_var = tk.StringVar()
        self.vendor_email_var = tk.StringVar()
        self.doc_date_var = tk.StringVar(value=datetime.date.today().strftime("%Y-%m-%d"))
        self.doc_validity_var = tk.StringVar(value="30 Days")
        self.quotation_no_var = tk.StringVar(value="OM-2024-001")
        self.vendor_code_var = tk.StringVar(value="")
        self.tc_spacing_var = tk.DoubleVar(value=1.5)
        self.made_by_var = tk.StringVar(value="Application Engineer")
        self.approved_by_var = tk.StringVar(value="Customer")

        self.client_name_var = tk.StringVar()
        self.client_addr_var = tk.StringVar()
        self.client_contact_var = tk.StringVar()
        self.client_designation_var = tk.StringVar()
        self.client_email_var = tk.StringVar()
        self.rfq_no_var = tk.StringVar()
        self.revision_no_var = tk.StringVar()
        self.client_phone_var = tk.StringVar()
        self.revised_date_var = tk.StringVar(value=datetime.date.today().strftime("%Y-%m-%d"))
        self.project_var = tk.StringVar()
        
        # Optional Rows
        self.header_l6_label_var = tk.StringVar(value="")
        self.header_l6_val_var = tk.StringVar(value="")
        self.header_r6_label_var = tk.StringVar(value="")
        self.header_r6_val_var = tk.StringVar(value="")
        self.doc_title_var = tk.StringVar(value="QUOTATION")
        self.doc_subtitle_var = tk.StringVar(value="")

        # Dynamic Header Rows
        self.header_rows = []
        self._init_standard_header_rows()

        # Item Inputs
        self.btn_add_text = tk.StringVar(value="ADD ITEM") 

        self._setup_styles()
        self._build_scrollable_gui()
        self.update_currency_symbol()

        self.current_db_id = None
        self.root.after(5000, self.auto_save_loop)

        # Bind closing event to stop background threads/loops
        self.root.protocol("WM_DELETE_WINDOW", self.on_closing)
        
        self.check_license()

    def on_closing(self):
        """Safely handle application shutdown"""
        try:
            if self.auto_save_timer:
                 self.root.after_cancel(self.auto_save_timer)
                 self.auto_save_timer = None
        except Exception:
            pass
        self.root.destroy()

    def check_license(self):
        app_data_dir = os.getenv('APPDATA')
        my_folder = os.path.join(app_data_dir, "ODM_Quotation_Gen")
        
        if not os.path.exists(my_folder):
            os.makedirs(my_folder)
            
        license_file = os.path.join(my_folder, "license.key")
        machine_id = self.get_machine_id()
        license_valid = False
        
        if os.path.exists(license_file):
            try:
                with open(license_file, "r") as f:
                    saved_key = f.read()
                    if self.validate_key(machine_id, saved_key):
                        license_valid = True
            except:
                pass
        
        if not license_valid:
            self.show_license_window(machine_id)
        else:
            self.check_user_login()

    def get_machine_id(self):
        try:
            cmd = 'wmic csproduct get uuid'
            uuid = str(subprocess.check_output(cmd).decode().split('\n')[1].strip())
            return uuid
        except:
            return "UNKNOWN-ID-ERROR"

    def validate_key(self, machine_id, input_key):
        SECRET_SALT = "ODM-ONLINE'S_QUOTATION_SYSTEM_2026_SECRET"
        raw_data = f"{machine_id.strip()}{SECRET_SALT}"
        hashed = hashlib.sha256(raw_data.encode()).hexdigest()
        generated_key = hashed[:16].upper()
        formatted_generated = f"{generated_key[0:4]}-{generated_key[4:8]}-{generated_key[8:12]}-{generated_key[12:16]}"
        return input_key.strip() == formatted_generated

    def show_license_window(self, machine_id):
        self.root.withdraw()
        lic_win = tk.Toplevel(self.root)
        lic_win.title("Product Activation")
        lic_win.geometry("800x800")
        lic_win.config(bg="#2c3e50")
        
        def close_app():
            self.root.destroy()
            import sys; sys.exit()
            
        lic_win.protocol("WM_DELETE_WINDOW", close_app)

        tk.Label(lic_win, text="🚫 Software Not Activated", font=("Segoe UI", 18, "bold"), bg="#2c3e50", fg="#e74c3c").pack(pady=(20,10))
        tk.Label(lic_win, text="Please send this Machine ID to Admin to get your Key:", bg="#2c3e50", fg="white", font=("Arial", 10)).pack()
        
        mid_ent = tk.Entry(lic_win, justify='center', font=("Consolas", 12, "bold"), width=40)
        mid_ent.pack(pady=10)
        mid_ent.insert(0, machine_id)
        mid_ent.config(state='readonly')
        
        def copy_id():
            lic_win.clipboard_clear()
            lic_win.clipboard_append(machine_id)
            messagebox.showinfo("Copied", "ID Copied!", parent=lic_win)
        
        tk.Button(lic_win, text="Copy ID", command=copy_id, bg="grey", fg="white", font=("Arial", 8)).pack(pady=2)

        tk.Label(lic_win, text="Enter License Key:", bg="#2c3e50", fg="white", font=("Arial", 10, "bold")).pack(pady=(20, 5))
        key_ent = tk.Entry(lic_win, justify='center', font=("Consolas", 12), width=30)
        key_ent.pack(pady=5)
        
        def activate():
            user_key = key_ent.get()
            if self.validate_key(machine_id, user_key):
                try:
                    app_data_dir = os.getenv('APPDATA')
                    license_path = os.path.join(app_data_dir, "ODM_Quotation_Gen", "license.key")
                    with open(license_path, "w") as f:
                        f.write(user_key)
                    messagebox.showinfo("Success", "Software Activated Successfully!", parent=lic_win)
                    lic_win.destroy()
                    self.check_user_login()
                except Exception as e:
                    messagebox.showerror("Error", f"Could not save license: {e}", parent=lic_win)
            else:
                messagebox.showerror("Error", "Invalid License Key!", parent=lic_win)

        tk.Button(lic_win, text="ACTIVATE NOW", command=activate, bg="#27ae60", fg="white", font=("Arial", 12, "bold"), width=20).pack(pady=20)

    def check_user_login(self):
        self.cursor.execute("SELECT COUNT(*) FROM users")
        count = self.cursor.fetchone()[0]
        if count == 0:
            self.perform_setup()
        else:
            self.perform_login(None)

    def perform_setup(self):
        self.root.withdraw()
        setup_win = tk.Toplevel(self.root)
        setup_win.title("First Time Setup - Create Profile")
        setup_win.state('zoomed') # Full Window
        
        def on_setup_close():
            # Check if any user exists, if so go back to login, else exit
            try:
                self.cursor.execute("SELECT COUNT(*) FROM users")
                if self.cursor.fetchone()[0] > 0:
                    setup_win.destroy()
                    self.perform_login(None) # Show login
                else:
                    sys.exit()
            except:
                sys.exit()

        setup_win.protocol("WM_DELETE_WINDOW", on_setup_close) 
        
        main_fr = ttk.Frame(setup_win, padding=40)
        main_fr.pack(fill='both', expand=True, anchor='center')

        # Title
        tk.Label(main_fr, text="Welcome! Let's Setup Your Profile", font=("Segoe UI", 24, "bold"), fg="#2c3e50").pack(pady=(20, 10))

        # Login button if account already exists
        def go_to_login():
             setup_win.destroy()
             self.perform_login(None)

        ttk.Button(main_fr, text="Already have an account? Login Here", command=go_to_login, bootstyle="info-outline", width=30).pack(pady=(0, 20))

        # Container for Columns
        cols_frame = ttk.Frame(main_fr)
        cols_frame.pack(fill='both', expand=True)

        # Left Column: Basic Info
        left_col = ttk.Labelframe(cols_frame, text="User Credentials", padding=15, bootstyle="info")
        left_col.pack(side='left', fill='both', expand=True, padx=10)

        ttk.Label(left_col, text="Full Name:").pack(anchor='w')
        name_ent = ttk.Entry(left_col, width=40); name_ent.pack(pady=5, fill='x')
        
        ttk.Label(left_col, text="Set Username:").pack(anchor='w')
        user_ent = ttk.Entry(left_col, width=40); user_ent.pack(pady=5, fill='x')
        
        ttk.Label(left_col, text="Set Password:").pack(anchor='w')
        pass_ent = ttk.Entry(left_col, width=40, show="*"); pass_ent.pack(pady=5, fill='x')

        # Profile Picture Section
        ttk.Label(left_col, text="Profile Picture:").pack(anchor='w', pady=(15, 5))
        
        pic_frame = ttk.Frame(left_col)
        pic_frame.pack(fill='x', pady=5)
        
        self.preview_img_lbl = ttk.Label(pic_frame, text="No Image", background="#ecf0f1", anchor="center", font=("Arial", 9))
        self.preview_img_lbl.pack(side='left', padx=(0,10))
        self.preview_img_lbl.config(width=15) # Fixed placeholder width

        self.setup_pic_path = None
        
        def choose_pic():
            p = filedialog.askopenfilename(parent=setup_win, filetypes=[("Images", "*.png;*.jpg;*.jpeg")])
            if p:
                self.setup_pic_path = p
                btn_pic.config(text="Change Image", bootstyle="warning-outline")
                
                # Show Preview
                try:
                    load = Image.open(p)
                    load = load.resize((80, 80), Image.Resampling.LANCZOS)
                    self.setup_preview_photo = ImageTk.PhotoImage(load)
                    self.preview_img_lbl.config(image=self.setup_preview_photo, text="")
                    btn_set_confirm.pack(side='left', padx=5) # Show Set Button
                    btn_set_confirm.config(state='normal', text="Set Image")
                except Exception as e:
                    messagebox.showerror("Error", f"Invalid Image: {e}")

        def confirm_pic():
             messagebox.showinfo("Image Set", "Profile Image Set Successfully!", parent=setup_win)
             # Visual feedback only since save happens at verify
             btn_set_confirm.config(state='disabled', text="✅ Set")

        btn_pic = ttk.Button(pic_frame, text="Choose Image...", command=choose_pic, bootstyle="secondary-outline")
        btn_pic.pack(side='left', fill='x', expand=True)

        btn_set_confirm = ttk.Button(pic_frame, text="Set Image", command=confirm_pic, bootstyle="success")
        # Note: We don't pack it immediately, only after selection to show 'Set' option

        # Right Column: Security Questions
        right_col = ttk.Labelframe(cols_frame, text="Security Recovery (Create 3 Questions)", padding=15, bootstyle="warning")
        right_col.pack(side='right', fill='both', expand=True, padx=10)

        # Q1
        ttk.Label(right_col, text="Question 1:").pack(anchor='w')
        q1_ent = ttk.Entry(right_col, width=40); q1_ent.pack(fill='x', pady=(0, 2))
        ttk.Label(right_col, text="Answer 1:", font=("Arial", 8, "italic"), foreground="grey").pack(anchor='w')
        a1_ent = ttk.Entry(right_col, width=40); a1_ent.pack(fill='x', pady=(0, 10))

        # Q2
        ttk.Label(right_col, text="Question 2:").pack(anchor='w')
        q2_ent = ttk.Entry(right_col, width=40); q2_ent.pack(fill='x', pady=(0, 2))
        ttk.Label(right_col, text="Answer 2:", font=("Arial", 8, "italic"), foreground="grey").pack(anchor='w')
        a2_ent = ttk.Entry(right_col, width=40); a2_ent.pack(fill='x', pady=(0, 10))

        # Q3
        ttk.Label(right_col, text="Question 3:").pack(anchor='w')
        q3_ent = ttk.Entry(right_col, width=40); q3_ent.pack(fill='x', pady=(0, 2))
        ttk.Label(right_col, text="Answer 3:", font=("Arial", 8, "italic"), foreground="grey").pack(anchor='w')
        a3_ent = ttk.Entry(right_col, width=40); a3_ent.pack(fill='x', pady=(0, 10))

        
        def save_user():
            # Strip extra spaces
            full_name = name_ent.get().strip()
            username = user_ent.get().strip()
            password = pass_ent.get().strip()
            
            # Security
            q1, a1 = q1_ent.get().strip(), a1_ent.get().strip()
            q2, a2 = q2_ent.get().strip(), a2_ent.get().strip()
            q3, a3 = q3_ent.get().strip(), a3_ent.get().strip()

            if not (full_name and username and password and q1 and a1 and q2 and a2 and q3 and a3):
                messagebox.showwarning("Input", "All fields (Credentials & 3 Q/A) are required!", parent=setup_win)
                return

            try:
                # Save to DB
                self.cursor.execute("""
                    INSERT INTO users (full_name, username, password, q1, a1, q2, a2, q3, a3, profile_pic_path, is_premium)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 0)
                """, (full_name, username, password, q1, a1, q2, a2, q3, a3, self.setup_pic_path))
                
                self.conn.commit()
                messagebox.showinfo("Success", "Profile Created Successfully! Login to continue.", parent=setup_win)
                setup_win.destroy()
                self.perform_login(None) 
            except Exception as e:
                messagebox.showerror("Error", f"Error saving profile: {e}", parent=setup_win)
                # If username exists error, handle it gracefully
                if "UNIQUE constraint" in str(e):
                    messagebox.showerror("Error", "Username already taken.", parent=setup_win)

        # Action Buttons
        btn_frame = ttk.Frame(main_fr)
        btn_frame.pack(pady=40)

        ttk.Button(btn_frame, text="SAVE & START SYSTEM", command=save_user, style="success.TButton", width=30).pack(pady=5, ipady=5)
        # ttk.Button(btn_frame, text="Already have an account",command=self.perform_login(None),style="info.TButton").pack(pady=5,ipady=5)
    # def perform_login(self, user_data):
    #     if user_data:
    #         self.root.deiconify()
    #         self.vendor_company_var.set(user_data[3]) 
    #         self.current_username = user_data[1]
    #         from dashboard import DashboardPanel
    #         self.current_dashboard = DashboardPanel(self)
    #         return 

    #     self.root.withdraw()
    #     login_win = tk.Toplevel(self.root)
    #     login_win.title("Login System")
    #     login_win.geometry("700x600")
    #     login_win.protocol("WM_DELETE_WINDOW", lambda: sys.exit())

    #     tk.Label(login_win, text="SECURE LOGIN", font=("Segoe UI", 18, "bold")).pack(pady=20)
        
    #     f = tk.Frame(login_win); f.pack(pady=5)
    #     tk.Label(f, text="Username:").pack(anchor='w')
    #     u_ent = ttk.Entry(f, width=30); u_ent.pack(pady=2)
        
    #     tk.Label(f, text="Password:").pack(anchor='w', pady=(10,0))
    #     p_ent = ttk.Entry(f, width=30, show="*"); p_ent.pack(pady=2)

    #     def try_login():
    #         username = u_ent.get().strip()
    #         password = p_ent.get().strip()
    #         self.cursor.execute("SELECT * FROM users WHERE username=? AND password=?", (username, password))
    #         row = self.cursor.fetchone()
    #         if row:
    #             login_win.destroy()
    #             self.root.deiconify()
    #             self.vendor_company_var.set(row[3]) 
    #             self.current_username = row[1]      
    #             from dashboard import DashboardPanel
    #             self.current_dashboard = DashboardPanel(self)
    #         else:
    #             messagebox.showerror("Failed", "Wrong Username or Password", parent=login_win)
    def perform_login(self, user_data=None):
        if user_data:
            self.root.deiconify()
            self.current_username = user_data[1]
            
            # Load Profile Pic Path if available
            try:
                # user_data is row from users table. schema changed so fetching by name is better, but user_data might be raw tuple
                # Let's re-fetch safely
                self.cursor.execute("SELECT profile_pic_path, theme_preference FROM users WHERE username=?", (self.current_username,))
                urow = self.cursor.fetchone()
                if urow:
                    self.current_user_pic_path = urow[0]  # Store for Dashboard
                    theme = urow[1] if urow[1] else "cosmo"
                    self.style = ThemeManager.apply_theme(self.root, theme)
            except Exception as e:
                print(f"Login data fetch error: {e}")
                self.current_user_pic_path = None

            from dashboard import DashboardPanel
            self.current_dashboard = DashboardPanel(self)
            return 

        self.root.withdraw()
        login_win = tk.Toplevel(self.root)
        login_win.title("ODM Secure Login")
        login_win.state('zoomed') # ZOOMED / MAXIMIZED
        login_win.protocol("WM_DELETE_WINDOW", lambda: sys.exit())

        # Center Frame
        main_fr = ttk.Frame(login_win, padding=30)
        main_fr.place(relx=0.5, rely=0.5, anchor='center')

        tk.Label(main_fr, text="SECURE LOGIN", font=("Segoe UI", 24, "bold"), fg="#2c3e50").pack(pady=(10, 20))
        
        ttk.Label(main_fr, text="Username:").pack(anchor='w')
        u_ent = ttk.Entry(main_fr, width=40, font=("Segoe UI", 11)); u_ent.pack(pady=5)
        
        ttk.Label(main_fr, text="Password:").pack(anchor='w', pady=(10, 0))
        p_ent = ttk.Entry(main_fr, width=40, show="*", font=("Segoe UI", 11)); p_ent.pack(pady=5)

        # ✅ REMEMBER ME LOGIC
        self.remember_me_var = tk.BooleanVar(value=False)
        rem_file = os.path.join(os.getenv('APPDATA'), "odm_remember.json")
        
        try:
            if os.path.exists(rem_file):
                with open(rem_file, 'r') as f:
                    saved = json.load(f)
                    u_ent.insert(0, saved.get('user', ''))
                    p_ent.insert(0, saved.get('pass', ''))
                    self.remember_me_var.set(True)
        except Exception as e:
            print(f"Error loading remember me data: {e}")
            # Optionally delete the corrupt file
            try: os.remove(rem_file)
            except: pass

        ttk.Checkbutton(main_fr, text="Remember Me", variable=self.remember_me_var, bootstyle="round-toggle").pack(anchor='w', pady=10)

        def try_login():
            username = u_ent.get().strip()
            password = p_ent.get().strip()
            self.cursor.execute("SELECT * FROM users WHERE username=? AND password=?", (username, password))
            row = self.cursor.fetchone()
            if row:
                try:
                    if self.remember_me_var.get():
                        with open(rem_file, 'w') as f: json.dump({'user': username, 'pass': password}, f)
                    elif os.path.exists(rem_file):
                        os.remove(rem_file)
                except Exception as e:
                    print(f"Error saving remember me data: {e}")

                login_win.destroy()
                self.perform_login(row)
            else:
                messagebox.showerror("Failed", "Invalid Credentials", parent=login_win)

        ttk.Button(main_fr, text="LOGIN", command=try_login, bootstyle="success", width=25).pack(pady=20)

        # ✅ RECOVERY Links
        ttk.Separator(main_fr, orient='horizontal').pack(fill='x', pady=15)
        
        ttk.Button(main_fr, text="Forgot Password?", command=lambda: self.forgot_password_flow(u_ent.get().strip(), login_win), 
                   bootstyle="link").pack(pady=5)

        ttk.Label(main_fr, text="Don't have a profile?", font=("Arial", 9), foreground="grey").pack(pady=(15, 0))
        ttk.Button(main_fr, text="Setup New Profile", command=lambda: [login_win.destroy(), self.perform_setup()], 
                   bootstyle="secondary-outline", width=25).pack(pady=5)

    def forgot_password_flow(self, prefill_user, parent_win):
        rec_win = tk.Toplevel(parent_win)
        rec_win.title("Password Recovery")
        rec_win.geometry("600x600")
        
        tk.Label(rec_win, text="Password Recovery", font=("Segoe UI", 16, "bold")).pack(pady=15)
        
        f = ttk.Frame(rec_win, padding=20)
        f.pack(fill='both', expand=True)

        tk.Label(f, text="Enter Username to fetch questions:").pack(anchor='w')
        u_e = ttk.Entry(f, width=40)
        u_e.pack(pady=5)
        if prefill_user: u_e.insert(0, prefill_user)

        # Container for Qs
        q_frame = ttk.Frame(f)
        q_frame.pack(fill='both', expand=True, pady=10)

        self.questions_loaded = False
        self.real_answers = []

        def load_questions():
            user = u_e.get().strip()
            if not user: return
            
            # Fetch q1, a1...
            try:
                self.cursor.execute("SELECT q1, a1, q2, a2, q3, a3 FROM users WHERE username=?", (user,))
                row = self.cursor.fetchone()
                
                for w in q_frame.winfo_children(): w.destroy()
                
                if row and row[0]: # Has Qs
                    self.questions_loaded = True
                    self.real_answers = [row[1], row[3], row[5]] # a1, a2, a3
                    
                    tk.Label(q_frame, text="Please answer the following:", font=("bold")).pack(pady=5)
                    
                    self.ans_entries = []
                    
                    # Q1
                    tk.Label(q_frame, text=f"Q1: {row[0]}").pack(anchor='w', pady=(5,0))
                    e1 = ttk.Entry(q_frame, width=40); e1.pack(pady=2)
                    self.ans_entries.append(e1)
                    
                    # Q2
                    tk.Label(q_frame, text=f"Q2: {row[2]}").pack(anchor='w', pady=(5,0))
                    e2 = ttk.Entry(q_frame, width=40); e2.pack(pady=2)
                    self.ans_entries.append(e2)
                    
                    # Q3
                    q3_text = row[4] if row[4] else "Question 3 (Not Set)"
                    tk.Label(q_frame, text=f"Q3: {q3_text}").pack(anchor='w', pady=(5,0))
                    e3 = ttk.Entry(q_frame, width=40); e3.pack(pady=2)
                    self.ans_entries.append(e3)
                    
                else:
                    tk.Label(q_frame, text="User not found or no questions set.", fg="red").pack(pady=20)
                    self.questions_loaded = False
            except Exception as e:
                print(e)

        ttk.Button(f, text="Load Questions", command=load_questions, style="info.TButton").pack(pady=5)
        
        # Auto-load if username keyin
        if prefill_user:
            load_questions()
        
        def verify_answers():
            if not self.questions_loaded: return
            
            username = u_e.get().strip()
            score = 0
            for i, ent in enumerate(self.ans_entries):
                user_ans = ent.get().strip().lower()
                real_ans = self.real_answers[i].strip().lower() if self.real_answers[i] else ""
                
                if user_ans and user_ans == real_ans:
                    score += 1
            
            if score >= 2:
                # Success
                new_pass = simpledialog.askstring("Reset Password", "Authentication Verified!\nEnter New Password:", parent=rec_win, show='*')
                if new_pass:
                    self.cursor.execute("UPDATE users SET password=? WHERE username=?", (new_pass, username))
                    self.conn.commit()
                    messagebox.showinfo("Success", "Password Changed! Please Login.", parent=rec_win)
                    rec_win.destroy()
            else:
                messagebox.showerror("Failed", f"Verification Failed.\nCorrect Answers: {score}/3\nNeed at least 2 correct.", parent=rec_win)
                # Show Re-create profile option
                btn_reset_profile.pack(pady=10)

        verify_btn = ttk.Button(f, text="Verify & Reset", command=verify_answers, style="success.TButton")
        verify_btn.pack(pady=20)

        # Fallback Option (Initially Hidden)
        def fallback_reset():
            if messagebox.askyesno("Setup New Profile", "This will overwrite the database keys for this user (if local DB). Continue?"):
                rec_win.destroy()
                parent_win.destroy()
                self.perform_setup()

        btn_reset_profile = ttk.Button(f, text="Forgot Answers? Setup New Profile", command=fallback_reset, style="danger.Outline.TButton")

    # =========================================================================
    # DATABASE METHODS
    # =========================================================================
    # def init_database(self):
    #     self.current_username = "admin"
    #     server_name = r'.\SQLEXPRESS'
    #     database_name = 'QuotationDB'
    #     try:
    #         conn_str = (f'DRIVER={{ODBC Driver 17 for SQL Server}};SERVER={server_name};DATABASE={database_name};Trusted_Connection=yes;TrustServerCertificate=yes;')
    #         self.conn = pyodbc.connect(conn_str)
    #         self.cursor = self.conn.cursor()
            
    #         # Tables Creation (Jo pehle tha waisa hi rahega)
    #         tables = [
    #             "quotations", "tax_invoices", "commercial_invoices", "delivery_challans"
    #         ]
            
    #         # --- MAGIC FIX: "created_by" column add karna ---
    #         for tbl in tables:
    #             # Table banao agar nahi hai
    #             self.cursor.execute(f"""
    #                 IF NOT EXISTS (SELECT * FROM sysobjects WHERE name='{tbl}' AND xtype='U')
    #                 CREATE TABLE {tbl} (
    #                     id INT IDENTITY(1,1) PRIMARY KEY,
    #                     ref_no NVARCHAR(50),
    #                     client_name NVARCHAR(255),
    #                     date NVARCHAR(50),
    #                     grand_total FLOAT,
    #                     full_data NVARCHAR(MAX),
    #                     created_by NVARCHAR(50) DEFAULT 'admin'  -- <--- YE NAYA COLUMN HAI
    #                 )
    #             """)
                
    #             # Agar table pehle se bana hai, to check karo 'created_by' column hai ya nahi
    #             try:
    #                 self.cursor.execute(f"SELECT created_by FROM {tbl} WHERE 1=0")
    #             except:
    #                 # Column nahi mila, to add kar do
    #                 self.conn.commit() # Transaction clear
    #                 print(f"Adding 'created_by' column to {tbl}...")
    #                 self.cursor.execute(f"ALTER TABLE {tbl} ADD created_by NVARCHAR(50) DEFAULT 'admin'")
    #                 self.conn.commit()

    #         # Users Table
    #         self.cursor.execute("""
    #             IF NOT EXISTS (SELECT * FROM sysobjects WHERE name='users' AND xtype='U')
    #             CREATE TABLE users (
    #                 id INT IDENTITY(1,1) PRIMARY KEY,
    #                 full_name NVARCHAR(100),
    #                 username NVARCHAR(50) UNIQUE,
    #                 password NVARCHAR(100),
    #                 is_premium INT DEFAULT 0
    #             )
    #         """)
    #         self.conn.commit()
    #         print(" Database Connected & Updated for Multi-User!")
            
    #     except Exception as e:
    #         messagebox.showerror("Database Error", f"SQL Server Connection Failed:\n{e}")
    def init_database(self):
        # ✅ Ab ye line sahi se indent hai
        if hasattr(self, 'conn') and self.conn:
            print("ℹ️ Database already initialized.")
            return

        self.db_name = "QuotationManager_Final.db"
        self.conn = sqlite3.connect(self.db_name)
        self.cursor = self.conn.cursor()

        # --- 1. BASIC TABLES SETUP ---
        tables_config = {
            "users": "id INTEGER PRIMARY KEY AUTOINCREMENT, username TEXT NOT NULL UNIQUE, password TEXT NOT NULL, full_name TEXT, role TEXT DEFAULT 'user', created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP",
            "quotations": "id INTEGER PRIMARY KEY AUTOINCREMENT, client_name TEXT, date TEXT, grand_total REAL, ref_no TEXT, created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP",
            "tax_invoices": "id INTEGER PRIMARY KEY AUTOINCREMENT, client_name TEXT, date TEXT, grand_total REAL, ref_no TEXT, created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP",
            "commercial_invoices": "id INTEGER PRIMARY KEY AUTOINCREMENT, client_name TEXT, date TEXT, grand_total REAL, ref_no TEXT, created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP",
            "delivery_challans": "id INTEGER PRIMARY KEY AUTOINCREMENT, client_name TEXT, date TEXT, grand_total REAL, ref_no TEXT, created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP"
        }

        for table, columns in tables_config.items():
            self.cursor.execute(f"CREATE TABLE IF NOT EXISTS {table} ({columns})")

        # --- 2. SMART COLUMNS UPDATER ---
        # Note: Ye function bhi indent hona chahiye kyunke ye init_database ke andar hai
        def add_column_if_missing(table, column, col_type):
            try:
                self.cursor.execute(f"SELECT {column} FROM {table} LIMIT 1")
            except sqlite3.OperationalError:
                print(f"⚙️ Updating DB: Adding '{column}' to '{table}'...")
                try:
                    self.cursor.execute(f"ALTER TABLE {table} ADD COLUMN {column} {col_type}")
                    self.conn.commit()
                except Exception as e:
                    print(f"Warning: Could not add {column} to {table}: {e}")

        # --- 3. APPLYING STANDARDS ---
        data_tables = ["quotations", "tax_invoices", "commercial_invoices", "delivery_challans"]
        
        for tbl in data_tables:
            add_column_if_missing(tbl, "ref_no", "TEXT")
            add_column_if_missing(tbl, "date", "TEXT")
            add_column_if_missing(tbl, "grand_total", "REAL")
            add_column_if_missing(tbl, "full_data", "TEXT")
            add_column_if_missing(tbl, "created_by", "TEXT")
            
            try:
                self.cursor.execute(f"CREATE UNIQUE INDEX IF NOT EXISTS idx_{tbl}_ref_no ON {tbl} (ref_no)")
            except:
                pass

        add_column_if_missing("users", "security_answer", "TEXT")
        add_column_if_missing("users", "is_premium", "INTEGER DEFAULT 0")
        add_column_if_missing("users", "theme_preference", "TEXT DEFAULT 'cosmo'")
        add_column_if_missing("users", "q1", "TEXT")
        add_column_if_missing("users", "a1", "TEXT")
        add_column_if_missing("users", "q2", "TEXT")
        add_column_if_missing("users", "a2", "TEXT")
        add_column_if_missing("users", "q3", "TEXT")
        add_column_if_missing("users", "a3", "TEXT")
        add_column_if_missing("users", "profile_pic_path", "TEXT")

        self.conn.commit()
        print("✅ Database Connected & Schema Synchronized (Errors Fixed)")
    def backup_database(self):
        """Creates a Universal ZIP Backup including Database and All Assets (Logos, Pics)"""
        try:
            timestamp = datetime.datetime.now().strftime("%Y-%m-%d_%H%M%S")
            file_path = filedialog.asksaveasfilename(
                defaultextension=".zip",
                filetypes=[("Quotation Universal Backup", "*.zip")],
                initialfile=f"Full_Backup_{timestamp}.zip"
            )
            
            if not file_path: return
            
            self.conn.commit()
            
            # Temporary directory create karein assets gather karne ke liye
            with tempfile.TemporaryDirectory() as tmpdir:
                # 1. Database copy karein
                db_copy = os.path.join(tmpdir, self.db_name)
                shutil.copy2(self.db_name, db_copy)
                
                # 2. Assets gather karein (Logos, Profile Pics etc)
                # Hum un files ko track kareinge jo external hain
                assets_dir = os.path.join(tmpdir, "assets")
                os.makedirs(assets_dir)
                
                # Sab paths ka map banayenge taake restore pe paths fix ho sakein
                # Note: Filhal hum poora database file hi backup kar rahay hain.
                # Agar user ne pictures computer ke kisi folder mein rakhi hain, to logic ye hai
                # ke hum unhe bhi bundle karein.
                
                # ZIP mein convert karein
                with zipfile.ZipFile(file_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                    # Database add karein
                    zipf.write(db_copy, arcname=self.db_name)
                    
                    # Optional: Yahan assets folder add ho sakta hai agar hum images auto-copy feature implement karein.
                    # Filhal, DB restore hi priority hai, lekin images handle karna robust banata hai.
                
                messagebox.showinfo("Success", f"Universal Backup Created!\n\nLocation: {file_path}")
                
        except Exception as e:
            messagebox.showerror("Backup Error", f"Failed: {e}")

    def restore_database(self):
        """Restores from a ZIP Universal Backup"""
        if not messagebox.askyesno("Confirm", "Restoring will OVERWRITE current data. Continue?"):
            return

        try:
            file_path = filedialog.askopenfilename(filetypes=[("Backup ZIP", "*.zip")])
            if not file_path: return

            with zipfile.ZipFile(file_path, 'r') as zipf:
                # Arcnames check karein
                if self.db_name not in zipf.namelist():
                    messagebox.showerror("Error", "Invalid Backup File: Database not found inside.")
                    return
                
                self.conn.close() # Close existing
                
                # Extract Database
                zipf.extract(self.db_name, path=".")
                
                # Re-init
                self.init_database()
                
            messagebox.showinfo("Success", "Restoration Complete!\nAll quotations, users and data recovered.")
            
        except Exception as e:
            messagebox.showerror("Restore Error", f"Failed: {e}")
            self.init_database()

    def auto_save_loop(self):
        """Ye function har 5 second baad chalega"""
        
        # FIX: Check if root exists to avoid "application has been destroyed" error
        try:
            if not self.root.winfo_exists():
                return
        except Exception:
            return

        # FIX: Sirf tab Auto-save karein agar ID majood ho (Yani file pehle se saved ho)
        # Agar nayi file hai, to hum wait karenge ke user pehli dafa khud 'Save' ka button dabaye.
        # Is se duplicates banna band ho jayenge.
        if self.current_db_id is not None and self.client_name_var.get().strip():
             self.save_to_database(silent=True)
            
        # Loop ko dobara schedule karein
        # Store the after ID so we can cancel it if needed (optional but good practice)
        self.auto_save_timer = self.root.after(5000, self.auto_save_loop)
        
    def save_to_database(self, silent=False):
        if not self.client_name_var.get().strip():
            if not silent: messagebox.showwarning("Error", "Client Name is required!", parent=self.root)
            return

        # Data Pack karna
        data_packet = {
            "header": {
                "quot_no": self.quotation_no_var.get(), "date": self.doc_date_var.get(),
                "validity": self.doc_validity_var.get(), "vendor_code": self.vendor_code_var.get(),
                "client_name": self.client_name_var.get(), "client_addr": self.client_addr_var.get(),
                "client_contact": self.client_contact_var.get(), "client_phone": self.client_phone_var.get(),
                "client_email": self.client_email_var.get(), "client_desig": self.client_designation_var.get(),
                "rfq": self.rfq_no_var.get(), "rev_date": self.revised_date_var.get(),
                "footer": self.footer_text_var.get(), "made_by": self.made_by_var.get(),
                "approved_by": self.approved_by_var.get()
            },
            "financial": {"curr": self.currency_var.get(), "gst": self.gst_rate_var.get()},
            "items": self.items_data,
            "colors": self.row_colors,
            "extra": [(k, v.get()) for k, v in self.extra_fields],
            "terms": self.terms_txt.get("1.0", "end-1c")
        }
        json_str = json.dumps(data_packet)

        # Total Calculation
        try:
            import re
            txt_total = self.total_lbl.cget("text")
            match = re.search(r"[\d,]+\.?\d*", txt_total)
            gt = float(match.group().replace(',', '')) if match else 0.0
        except: gt = 0.0

        try:
            # --- 1. UPDATE LOGIC (Record pehle se load hai) ---
            if self.current_db_id:
                self.cursor.execute("""
                    UPDATE quotations 
                    SET ref_no=?, client_name=?, date=?, grand_total=?, full_data=?, created_by=? 
                    WHERE id=?
                """, (self.quotation_no_var.get(), self.client_name_var.get(), self.doc_date_var.get(), gt, json_str, self.current_username, self.current_db_id))
            
            # --- 2. INSERT LOGIC (Naya record hai) ---
            else:
                # UNIQUE CHECK: Pehle check karein ke ye Ref No pehle se DB mein to nahi?
                self.cursor.execute("SELECT id FROM quotations WHERE ref_no=?", (self.quotation_no_var.get(),))
                already_exists = self.cursor.fetchone()
                
                if already_exists:
                    if not silent: 
                        messagebox.showerror("Duplicate Error", f"Quotation No '{self.quotation_no_var.get()}' Already saved!\n\nDelete the number or edit the previous entry.")
                    return # Save cancel kar dein

                # Agar duplicate nahi hai, to Insert karein
                self.cursor.execute("""
                    INSERT INTO quotations (ref_no, client_name, date, grand_total, full_data, created_by) 
                    VALUES (?,?,?,?,?,?)
                """, (self.quotation_no_var.get(), self.client_name_var.get(), self.doc_date_var.get(), gt, json_str, self.current_username))
                
                self.current_db_id = self.cursor.lastrowid 
                print(f"New ID Assigned: {self.current_db_id}")

            self.conn.commit()
            
            if not silent:
                toast = ToastNotification(
                    title="Saved Successfully",
                    message=f"Quotation #{self.quotation_no_var.get()} has been saved.",
                    duration=3000,
                    bootstyle="success",
                    position=(50, 50, 'ne')
                )
                toast.show_toast()

        except Exception as e:
            if not silent: messagebox.showerror("DB Error", f"Database error: {str(e)}")

    def update_app_theme(self, event=None):
        """
        Callback used by the Dashboard theme combobox.

        Reads the selected theme from the dashboard, applies it globally,
        and persists the preference to the current user's record.
        """
        # Dashboard / combobox may not yet be initialized
        dash = getattr(self, "current_dashboard", None)
        if dash is None or not hasattr(dash, "theme_var"):
            return

        theme_name = dash.theme_var.get().strip()
        if not theme_name:
            return

        # Apply the theme across the entire app
        self.style = ThemeManager.apply_theme(self.root, theme_name)

        # Persist the preference for this user
        try:
            if hasattr(self, "cursor") and hasattr(self, "conn") and getattr(self, "current_username", None):
                self.cursor.execute(
                    "UPDATE users SET theme_preference=? WHERE username=?",
                    (theme_name, self.current_username),
                )
                self.conn.commit()
        except Exception as e:
            print(f"Theme preference save error: {e}")
    # --- HELPER: RESET FORM (NEW FILE) ---
    def reset_form(self):
        # 1. ID Khatam karein (Taake nayi file bane)
        self.current_db_id = None
        
        # 2. Header Variables Khali karein
        self.client_name_var.set("")
        self.client_addr_var.set("")
        self.client_contact_var.set("")
        self.client_phone_var.set("")
        self.client_email_var.set("")
        self.client_designation_var.set("")
        self.rfq_no_var.set("")
        
        # Date waghaira ko default par layen
        self.quotation_no_var.set("OM-2024-NEW") 
        self.doc_date_var.set(datetime.date.today().strftime("%Y-%m-%d"))
        
        # 3. Items aur Colors saaf karein
        self.items_data = []
        self.row_colors = {}
        
        # 4. Extra Fields saaf karein
        for w in self.extra_cont.winfo_children(): w.destroy()
        self.extra_fields = []

        # 5. Screen Refresh
        self.refresh_tree()
        self.recalc_all()
        
        # 6. Title Update
        self.root.title("Professional Quotation Generator - [NEW FILE]")

    def load_history_window(self):
        # --- 1. WINDOW SETUP ---
        win = tk.Toplevel(self.root)
        win.title("Quotation History & Manager")
        win.geometry("1150x650")
        win.transient(self.root)
        win.grab_set()

        # --- 2. TREEVIEW SETUP ---

        # --- 2. TREEVIEW SETUP ---
        cols = ("chk", "ID", "Ref No", "Client", "Date", "Amount")
        tree = ttk.Treeview(win, columns=cols, show='headings', selectmode='browse')
        
        tree.heading("chk", text="✔");       tree.column("chk", width=40, anchor="center")
        tree.heading("ID", text="ID");       tree.column("ID", width=60, anchor="center")
        tree.heading("Ref No", text="Ref No"); tree.column("Ref No", width=120)
        tree.heading("Client", text="Client Name"); tree.column("Client", width=300)
        tree.heading("Date", text="Date");   tree.column("Date", width=100)
        tree.heading("Amount", text="Total"); tree.column("Amount", width=120)
        
        sb = ttk.Scrollbar(win, orient="vertical", command=tree.yview)
        tree.configure(yscrollcommand=sb.set)
        sb.pack(side='right', fill='y')
        tree.pack(fill='both', expand=True, padx=10, pady=10)

        # --- 3. HELPER FUNCTIONS ---
        def get_selected_db_id():
            sel = tree.selection()
            if not sel: return None
            vals = tree.item(sel[0])['values']
            try:
                clean = ''.join(filter(str.isdigit, str(vals[1])))
                return int(clean) if clean else None
            except: return None

        def refresh_list():
            for i in tree.get_children(): tree.delete(i)
            try:
                self.cursor.execute("SELECT id, ref_no, client_name, date, grand_total FROM quotations ORDER BY id DESC")
                rows = self.cursor.fetchall()
                for row in rows:
                    tree.insert("", "end", values=("☐", row[0], row[1], row[2], row[3], row[4]))
            except Exception as e:
                messagebox.showerror("History Error", f"Could not load data:\n{e}")

        def on_click(event):
            region = tree.identify("region", event.x, event.y)
            if region == "cell":
                col = tree.identify_column(event.x)
                row_id = tree.identify_row(event.y)
                if col == "#1" and row_id: 
                    vals = tree.item(row_id)['values']
                    new_mark = "☑" if str(vals[0]) == "☐" else "☐"
                    tree.item(row_id, values=(new_mark,) + tuple(vals[1:]))
        tree.bind("<Button-1>", on_click)

        # --- 4. ACTION: LOAD QUOTATION (CRITICAL FIX HERE) ---
        def load_quotation(event=None):
            db_id = get_selected_db_id()
            if not db_id: 
                messagebox.showwarning("Load", "Please select a quotation first.")
                return

            try:
                self.cursor.execute("SELECT full_data FROM quotations WHERE id=?", (db_id,))
                row = self.cursor.fetchone()
                if row:
                    win.destroy()
 
                    if self.current_dashboard:
                        try:
                            # Koshish karein destroy karne ki
                            if hasattr(self.current_dashboard, 'destroy'):
                                self.current_dashboard.destroy()
                            elif hasattr(self.current_dashboard, 'frame'):
                                self.current_dashboard.frame.destroy()
                        except: pass
                        self.current_dashboard = None

                    # 3. Data Load Karein
                    print(f"Loading ID: {db_id}")
                    self.restore_data(row[0])
                    self.current_db_id = db_id
                    
                    # 4. Main Window ko Upar Layen
                    self.root.deiconify() # Agar minimize thi to restore karein
                    self.root.lift()      # Screen ke sab se upar layen
                    self.root.title(f"Quotation Generator - [Loaded ID: {db_id}]")
                    self.root.update()    # Screen Refresh Force Karein
                    
                    messagebox.showinfo("Success", "Quotation Loaded Successfully!")
                else:
                    messagebox.showerror("Error", "Data not found in DB.")
            except Exception as e:
                messagebox.showerror("Load Error", f"Failed:\n{e}")


        def rename_record():
            sel = tree.selection()
            if not sel:
                messagebox.showwarning("Rename", "Please select a row first.")
                return
            
            db_id = get_selected_db_id()
            current_ref = tree.item(sel[0])['values'][2]

            new_name = simpledialog.askstring("Rename", "Enter New Ref No:", initialvalue=current_ref, parent=win)
            if new_name and db_id:
                try:
                    self.cursor.execute("UPDATE quotations SET ref_no=? WHERE id=?", (new_name, db_id))
                    
                    # Update JSON inside DB too
                    self.cursor.execute("SELECT full_data FROM quotations WHERE id=?", (db_id,))
                    row = self.cursor.fetchone()
                    if row:
                        import json
                        d = json.loads(row[0])
                        d['header']['quot_no'] = new_name
                        self.cursor.execute("UPDATE quotations SET full_data=? WHERE id=?", (json.dumps(d), db_id))
                    
                    self.conn.commit()
                    
                    # Update List
                    vals = list(tree.item(sel[0])['values'])
                    vals[2] = new_name
                    tree.item(sel[0], values=vals)
                    messagebox.showinfo("Success", "Renamed!")
                except Exception as e:
                    messagebox.showerror("Error", str(e))

        # --- 6. ACTION: CONVERT & DELETE ---
        def run_conversion(module_name, class_name, parent_win):
            db_id = get_selected_db_id()
            if not db_id:
                messagebox.showwarning("Convert", "Select a quotation.")
                return
            try:
                self.cursor.execute("SELECT full_data FROM quotations WHERE id=?", (db_id,))
                row = self.cursor.fetchone()
                if row:
                    parent_win.grab_release() 
                    parent_win.destroy()      
                    
                    import importlib
                    mod = importlib.import_module(module_name)
                    app_class = getattr(mod, class_name)
                    
                    new_win = tk.Toplevel(self.root)
                    app_instance = app_class(new_win, from_quotation_data=row[0])
                    new_win.state('zoomed')
                    new_win.lift()
                    new_win.focus_force()
                    new_win.attributes('-topmost', True)
                    new_win.after(1000, lambda: new_win.attributes('-topmost', False))
            except Exception as e:
                messagebox.showerror("Error", str(e))
        


        # Button callback functions
        def to_tax(): run_conversion("invoice", "InvoiceApp", win)
        def to_comm(): run_conversion("commercial", "CommercialApp", win)
        def to_dc(): run_conversion("delivery_challan", "DeliveryChallanApp", win)

        def delete_selected():
            targets = [c for c in tree.get_children() if str(tree.item(c)['values'][0]) == "☑"]
            if not targets and tree.selection(): targets.append(tree.selection()[0])
            
            if not targets:
                messagebox.showwarning("Delete", "Select items to delete.")
                return

            if not messagebox.askyesno("Confirm", f"Delete {len(targets)} items?"): return

            count = 0
            for item in targets:
                try:
                    vals = tree.item(item)['values']
                    clean = ''.join(filter(str.isdigit, str(vals[1])))
                    if clean:
                        self.cursor.execute("DELETE FROM quotations WHERE id=?", (int(clean),))
                        tree.delete(item)
                        count += 1
                except: pass
            self.conn.commit()
            if count: messagebox.showinfo("Success", f"Deleted {count} items.")

        # --- 7. BUTTON LAYOUT ---
        act_frame = ttk.LabelFrame(win, text="Create Documents")
        act_frame.pack(fill='x', padx=10, pady=5)
        ttk.Button(act_frame, text="📑 SalesTax Invoice", command=to_tax).pack(side='left', padx=5, pady=5, expand=True, fill='x')
        ttk.Button(act_frame, text="📄 Commercial Inv", command=to_comm).pack(side='left', padx=5, pady=5, expand=True, fill='x')
        ttk.Button(act_frame, text="🚚 Delivery Challan", command=to_dc).pack(side='left', padx=5, pady=5, expand=True, fill='x')

        mng_frame = ttk.Frame(win)
        mng_frame.pack(fill='x', padx=10, pady=10)

        # Load & Rename
        ttk.Button(mng_frame, text="📂 Load Quotation", command=load_quotation).pack(side='left', padx=5)
        ttk.Button(mng_frame, text="✏ Rename", command=rename_record).pack(side='left', padx=5)

        # Delete & Refresh
        del_btn = tk.Button(mng_frame, text="🗑 Delete Checked", bg="#ffcccc", fg="red", font=("Arial", 10, "bold"), command=delete_selected)
        del_btn.pack(side='right', padx=5)
        ttk.Button(mng_frame, text="🔄 Refresh List", command=refresh_list).pack(side='right', padx=5)

        tree.bind("<Double-1>", load_quotation)
        refresh_list()

    
        # def refresh_list():
        #     for i in tree.get_children():
        #         tree.delete(i)
          
        #     self.cursor.execute("SELECT id, ref_no, client_name, date, grand_total FROM quotations ORDER BY id DESC")
        #     rows = self.cursor.fetchall()
        #     for row in rows:
        #         tree.insert("", "end", values=row)


        # refresh_list()
    def go_to_dashboard(self):
    
        from dashboard import DashboardPanel   # <--- Ye line add karein
        self.current_dashboard = DashboardPanel(self)

    def restore_data(self, json_str):
        try:
            print("Restoring Data...") # Debugging
            data = json.loads(json_str)
            
            # 1. Restore Headers
            h = data.get("header", {})
            self.quotation_no_var.set(h.get("quot_no", ""))
            self.client_name_var.set(h.get("client_name", ""))
            self.doc_date_var.set(h.get("date", ""))
            self.doc_validity_var.set(h.get("validity", ""))
            self.vendor_code_var.set(h.get("vendor_code", ""))
            
            self.client_addr_var.set(h.get("client_addr", ""))
            self.client_contact_var.set(h.get("client_contact", ""))
            self.client_phone_var.set(h.get("client_phone", ""))
            self.client_email_var.set(h.get("client_email", ""))
            self.client_designation_var.set(h.get("client_desig", ""))
            
            self.rfq_no_var.set(h.get("rfq", ""))
            self.revised_date_var.set(h.get("rev_date", ""))
            self.made_by_var.set(h.get("made_by", ""))
            self.approved_by_var.set(h.get("approved_by", ""))
            self.footer_text_var.set(h.get("footer", ""))
            
            # 2. Restore Financials
            fin = data.get("financial", {})
            self.currency_var.set(fin.get("curr", "PKR"))
            self.gst_rate_var.set(fin.get("gst", 18.0))
            
            # 3. Restore Items (Clean Logic)
            self.items_data = [] 
            raw_items = data.get("items", [])
            for item in raw_items:
                clean = item.copy()
                # Numbers fix karein
                if 'price' in clean: clean['price'] = float(str(clean['price']).replace(',', ''))
                if 'qty' in clean: clean['qty'] = float(str(clean['qty']).replace(',', ''))
                self.items_data.append(clean)

            # 4. Restore Colors
            raw_colors = data.get("colors", {})
            self.row_colors = {int(k): v for k, v in raw_colors.items()}
            
            # 5. Restore Terms
            self.terms_txt.delete("1.0", "end")
            self.terms_txt.insert("1.0", data.get("terms", ""))
            
            # 6. Restore Extra Fields
            for w in self.extra_cont.winfo_children(): w.destroy()
            self.extra_fields = []
            for name, val in data.get("extra", []):
                fr = ttk.Frame(self.extra_cont); fr.pack(fill='x', pady=1)
                ttk.Label(fr, text=f"{name}:").pack(side='left', padx=5)
                var = tk.StringVar(value=val); ttk.Entry(fr, textvariable=var).pack(side='left', fill='x', expand=True)
                self.extra_fields.append((name, var))

            # --- CRITICAL UI UPDATES ---
            self.update_currency_symbol()
            self.refresh_tree()  # Ye line items ko screen par layegi
            self.recalc_all()    # Ye total update karegi
            
            # Ye command screen ko foran repaint karti hai
            self.root.update_idletasks()
            print("Data Restoration Complete.")

        except Exception as e:
            messagebox.showerror("Restore Error", f"Data load nahi hua:\n{e}")
    # =========================================================================
    # UI BUILDING & LAYOUT
    # =========================================================================
    def _setup_styles(self):
        # --- CUSTOM STYLES FOR MODERN LOOK ---
        style = self.style # ttkbootstrap style object
        
        # Fonts
        style.configure('.', font=('Segoe UI', 10))
        style.configure('Head.TLabel', font=('Segoe UI', 12, 'bold'), foreground="#3498db") # Blue Headings
        style.configure('BigTotal.TLabel', font=('Segoe UI', 24, 'bold'), foreground="#2ecc71") # Green Total
        
        # Treeview (Table) Styling
        style.configure("Treeview", 
                        background="#2c3e50", 
                        foreground="white", 
                        fieldbackground="#2c3e50", 
                        rowheight=30,
                        font=('Segoe UI', 10))
        
        style.configure("Treeview.Heading", 
                        font=('Segoe UI', 10, 'bold'), 
                        background="#34495e", 
                        foreground="white")
        
        # Card-like Frames
        style.configure('Card.TLabelframe', relief='solid', borderwidth=1)
        style.configure('Card.TLabelframe.Label', font=('Segoe UI', 11, 'bold'), foreground="#f39c12")

    def _build_scrollable_gui(self):
        canvas = tk.Canvas(self.root)
        scrollbar = ttk.Scrollbar(self.root, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)
        scrollable_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.configure(yscrollcommand=scrollbar.set, yscrollincrement=5)

        def _on_mousewheel(event): 
            try: 
                speed = int(-1 * (event.delta / 60)) 
                canvas.yview_scroll(speed, "units")        
            except: pass

        window_id = canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        def _conf(e):
            canvas.itemconfig(window_id, width=e.width)
        canvas.bind('<Configure>', _conf)

        canvas.configure(yscrollcommand=scrollbar.set)
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        self.root.bind_all("<MouseWheel>", _on_mousewheel)

        self._build_content(scrollable_frame)

    def _build_content(self, parent):
        main_container = ttk.Frame(parent) 
        main_container.pack(fill='both', expand=True, padx=15, pady=(15, 50))
        self._build_header_section(main_container)
        self._build_items_section(main_container)
        self._build_bottom_section(main_container)

    # def _build_header_section(self, parent):
    #     box = ttk.LabelFrame(parent, text="Document Header Setup")
    #     box.pack(fill='x', pady=5)
        
    #     btn_fr = ttk.Frame(box)
    #     btn_fr.pack(fill='x', pady=5)
    #     ttk.Button(btn_fr, text=" Save to Database", command=self.save_to_database).pack(side='left', padx=5)
    #     ttk.Button(btn_fr, text=" History", command=self.load_history_window).pack(side='left', padx=5)
    #     ttk.Button(btn_fr, text=" Dashboard", command=self.go_to_dashboard).pack(side='left', padx=5) 

    #     top_layout = ttk.Frame(box)
    #     top_layout.pack(fill='x', pady=(0, 10))
        
    #     v_logo_fr = ttk.Frame(top_layout)
    #     v_logo_fr.pack(side='left', anchor='nw')
        
    #     v_logo_row = ttk.Frame(v_logo_fr)
    #     v_logo_row.pack(fill='x', pady=2)
    #     ttk.Button(v_logo_row, text="+ Logo", width=8, command=lambda: self.load_logo("header")).pack(side='left')
    #     ttk.Button(v_logo_row, text="X", width=3, command=lambda: self.clear_logo("header")).pack(side='left', padx=2)
    #     ttk.Scale(v_logo_row, variable=self.v_logo_size_var, from_=0.5, to=7.5, length=80).pack(side='left', padx=5)
        
    #     self.header_logo_lbl = ttk.Label(v_logo_fr, text="(No Logo Selected)", foreground="grey")
    #     self.header_logo_lbl.pack(side='top', padx=5, anchor='w')

    #     m_logo_fr = ttk.Frame(top_layout)
    #     m_logo_fr.pack(side='left', expand=True, fill='x')
    #     m_inner = ttk.Frame(m_logo_fr)
    #     m_inner.pack(anchor='center')
        
    #     m_logo_row = ttk.Frame(m_inner)
    #     m_logo_row.pack(fill='x', pady=2)
    #     ttk.Button(m_logo_row, text="+ Logo (Mid)", width=10, command=lambda: self.load_logo("middle")).pack(side='left')
    #     ttk.Button(m_logo_row, text="X", width=3, command=lambda: self.clear_logo("middle")).pack(side='left', padx=2)
    #     ttk.Scale(m_logo_row, variable=self.m_logo_size_var, from_=0.5, to=3.5, length=60).pack(side='left', padx=5)
        
    #     self.middle_logo_lbl = ttk.Label(m_inner, text="(No Middle Logo)", foreground="grey")
    #     self.middle_logo_lbl.pack(side='top', padx=5)

    #     q_info_fr = ttk.Frame(top_layout)
    #     q_info_fr.pack(side='right', anchor='ne')
        
    #     title_font = ("Segoe Script", 28, "bold")
    #     ttk.Label(q_info_fr, text="Quotation", font=title_font, foreground="#2c3e50").pack(side='top', anchor='e')
        
    #     q_no_row = ttk.Frame(q_info_fr)
    #     q_no_row.pack(side='top', anchor='e', pady=5)
    #     ttk.Label(q_no_row, text="Quot No:", font=("Segoe UI", 10, "bold")).pack(side='left')
    #     ttk.Entry(q_no_row, textvariable=self.quotation_no_var, width=20).pack(side='left', padx=5)

    #     code_fr = ttk.Frame(box)
    #     code_fr.pack(fill='x', padx=10, pady=10)
        
    #     code_inner = ttk.Frame(code_fr)
    #     code_inner.pack(anchor='center')
    #     ttk.Label(code_inner, text="Vendor's Code:", font=("Segoe UI", 10, "bold")).pack(side='left')
    #     ttk.Entry(code_inner, textvariable=self.vendor_code_var, width=25, justify='center').pack(side='left', padx=10)

    #     self.header_grid_fr = ttk.Frame(box)
    #     self.header_grid_fr.pack(fill='x', padx=10, pady=10)
        
    #     self.refresh_custom_header_ui()

    #     bot_fr = ttk.Frame(box)
    #     bot_fr.pack(fill='x', pady=5)
    #     ttk.Button(bot_fr, text="⚙ Manage Rows", command=self.open_header_manager).pack(side='right', padx=10)
    #     ttk.Button(bot_fr, text="🎨 Set All Colors", command=self.pick_header_color).pack(side='right')
    def _build_header_section(self, parent):
        # 1. Main Card Container (Dashboard Style)
        # ttk.LabelFrame ki jagah hum style.create_card use karenge taake Dark color mile
        try:
            box = style.create_card(parent)
        except:
            # Fallback agar style import na ho
            box = ttk.Frame(parent)
            
        box.pack(fill='x', pady=10)

        # Title for the Card (Manually added because Frame doesn't have text attribute)
        title_fr = tk.Frame(box, bg=style.BG_CARD)
        title_fr.pack(fill='x', padx=15, pady=(10, 5))
        tk.Label(title_fr, text="DOCUMENT HEADER SETUP", font=style.FONT_H2, bg=style.BG_CARD, fg=style.ACCENT).pack(side='left')

        # --- TOP TOOLBAR ---
        # Background color set kiya taake Grey na dikhe
        btn_fr = tk.Frame(box, bg=style.BG_CARD)
        btn_fr.pack(fill='x', padx=15, pady=5)
        
        # Buttons with 'success' (Green) and 'info' styles
        ttk.Button(btn_fr, text="💾 Save to Database", bootstyle="success", command=self.save_to_database).pack(side='left', padx=2)
        ttk.Button(btn_fr, text="📂 History", bootstyle="info-outline", command=self.load_history_window).pack(side='left', padx=2)
        ttk.Button(btn_fr, text="🏠 Dashboard", bootstyle="primary-outline", command=self.go_to_dashboard).pack(side='left', padx=2)

        # --- LOGO & INFO SECTION ---
        top_layout = tk.Frame(box, bg=style.BG_CARD)
        top_layout.pack(fill='x', padx=15, pady=15)
        
        # 1. Left Logos
        v_logo_fr = tk.Frame(top_layout, bg=style.BG_CARD)
        v_logo_fr.pack(side='left', anchor='nw')
        
        # Header Logo Row
        v_logo_row = tk.Frame(v_logo_fr, bg=style.BG_CARD)
        v_logo_row.pack(fill='x', pady=2)
        # Small styled buttons
        ttk.Button(v_logo_row, text="+ Logo", bootstyle="success-link", width=6, command=lambda: self.load_logo("header")).pack(side='left')
        ttk.Button(v_logo_row, text="x", bootstyle="danger-link", width=2, command=lambda: self.clear_logo("header")).pack(side='left')
        ttk.Scale(v_logo_row, variable=self.v_logo_size_var, from_=0.5, to=7.5, bootstyle="success").pack(side='left', padx=5)
        
        self.header_logo_lbl = tk.Label(v_logo_fr, text="(No Logo Selected)", bg=style.BG_CARD, fg="grey", font=("Arial", 8))
        self.header_logo_lbl.pack(side='top', padx=5, anchor='w')

        # Mid Logo Row
        m_logo_fr = tk.Frame(top_layout, bg=style.BG_CARD)
        m_logo_fr.pack(side='left', expand=True, fill='x', padx=20)
        m_inner = tk.Frame(m_logo_fr, bg=style.BG_CARD)
        m_inner.pack(anchor='center')
        
        m_logo_row = tk.Frame(m_inner, bg=style.BG_CARD)
        m_logo_row.pack(fill='x', pady=2)
        ttk.Button(m_logo_row, text="+ Mid Logo", bootstyle="success-link", width=10, command=lambda: self.load_logo("middle")).pack(side='left')
        ttk.Button(m_logo_row, text="x", bootstyle="danger-link", width=2, command=lambda: self.clear_logo("middle")).pack(side='left')
        ttk.Scale(m_logo_row, variable=self.m_logo_size_var, from_=0.5, to=3.5, bootstyle="success").pack(side='left', padx=5)
        
        self.middle_logo_lbl = tk.Label(m_inner, text="(No Middle Logo)", bg=style.BG_CARD, fg="grey", font=("Arial", 8))
        self.middle_logo_lbl.pack(side='top', padx=5)

        # 2. Right Info (Title & Code)
        q_info_fr = tk.Frame(top_layout, bg=style.BG_CARD)
        q_info_fr.pack(side='right', anchor='ne')
        
        # Quotation Title (Green Accent)
        title_font = ("Segoe Script", 28, "bold")
        tk.Label(q_info_fr, text="Quotation", font=title_font, bg=style.BG_CARD, fg=style.ACCENT).pack(side='top', anchor='e')
        
        # Quotation No
        q_no_row = tk.Frame(q_info_fr, bg=style.BG_CARD)
        q_no_row.pack(side='top', anchor='e', pady=5)
        tk.Label(q_no_row, text="Quot No:", font=("Segoe UI", 10, "bold"), bg=style.BG_CARD, fg="white").pack(side='left')
        ttk.Entry(q_no_row, textvariable=self.quotation_no_var, width=20, bootstyle="success").pack(side='left', padx=5)

        # Vendor Code
        code_fr = tk.Frame(box, bg=style.BG_CARD)
        code_fr.pack(fill='x', padx=15, pady=5)
        code_inner = tk.Frame(code_fr, bg=style.BG_CARD)
        code_inner.pack(anchor='center')
        tk.Label(code_inner, text="Vendor's Code:", font=("Segoe UI", 10, "bold"), bg=style.BG_CARD, fg="white").pack(side='left')
        ttk.Entry(code_inner, textvariable=self.vendor_code_var, width=25, justify='center', bootstyle="secondary").pack(side='left', padx=10)

        # --- CUSTOM HEADER GRID ---
        # Container ka background change kiya
        self.header_grid_fr = tk.Frame(box, bg=style.BG_CARD) 
        self.header_grid_fr.pack(fill='x', padx=15, pady=10)
        
        self.refresh_custom_header_ui()

        # Bottom Buttons
        bot_fr = tk.Frame(box, bg=style.BG_CARD)
        bot_fr.pack(fill='x', pady=10, padx=15)
        ttk.Button(bot_fr, text="⚙ Manage Rows", bootstyle="secondary-outline", command=self.open_header_manager).pack(side='right', padx=10)
        ttk.Button(bot_fr, text="🎨 Set All Colors", bootstyle="warning-outline", command=self.pick_header_color).pack(side='right')

    def _build_items_section(self, parent):
                # 1. Items section ka main container banayein
                self.items_section_frame = ttk.LabelFrame(parent, text="Line Items")
                self.items_section_frame.pack(fill='both', expand=True, pady=10)
                
                # 2. Toolbar (Manage Columns aur Visibility buttons ke liye)
                tool_fr = ttk.Frame(self.items_section_frame)
                tool_fr.pack(fill='x', pady=5)
                
                ttk.Button(tool_fr, text="📊 Manage Columns", bootstyle="info-outline", command=self.open_column_manager).pack(side='left', padx=5)
                
                # Column visibility buttons build karein
                self._build_visibility_toolbar(tool_fr)

                # 3. Content build karein (Treeview aur Inputs)
                self._build_items_section_content(self.items_section_frame)
       

    # def _build_bottom_section(self, parent):
    #     # Main Container for Bottom
    #     bot_box = ttk.Frame(parent)
    #     bot_box.pack(fill='both', expand=True, pady=10)
        
    #     # 3 Columns Layout
    #     bot_box.columnconfigure(0, weight=1)
    #     bot_box.columnconfigure(1, weight=1)
    #     bot_box.columnconfigure(2, weight=1)

    #     # --- COL 1: FINANCIALS ---
    #     col1 = ttk.Frame(bot_box)
    #     col1.grid(row=0, column=0, sticky="nsew", padx=10)
        
    #     fin_fr = ttk.LabelFrame(col1, text="Financial Settings")
    #     fin_fr.pack(fill='x')
        
    #     r1 = ttk.Frame(fin_fr); r1.pack(fill='x', pady=5)
    #     ttk.Label(r1, text="Currency:").pack(side='left')
    #     ttk.Combobox(r1, textvariable=self.currency_var, values=["PKR", "USD", "EUR", "GBP"], width=5).pack(side='left', padx=5)
    #     self.currency_var.trace('w', self.update_currency_symbol)
        
    #     ttk.Label(r1, text="GST %:").pack(side='left', padx=10)
    #     gst_e = ttk.Entry(r1, textvariable=self.gst_rate_var, width=5)
    #     gst_e.pack(side='left')
    #     gst_e.bind('<FocusOut>', lambda e: self.recalc_all())

    #     self.total_lbl = ttk.Label(col1, text="Grand Total: 0.00", font=("Segoe UI", 16, "bold"), foreground="#2ecc71")
    #     self.total_lbl.pack(pady=15, anchor='w')

    #     # Extra Fields
    #     ttk.Button(col1, text="+ Add Extra Field", command=self.add_extra_field).pack(anchor='w')
    #     self.extra_cont = ttk.Frame(col1)
    #     self.extra_cont.pack(fill='x', pady=5)

    #     # --- COL 2: TERMS (Fixed Error Here) ---
    #     # Padding hata diya hai taake crash na ho
    #     lt_fr = ttk.LabelFrame(bot_box, text="Terms & Conditions") 
    #     lt_fr.grid(row=0, column=1, sticky="nsew", padx=10)
        
    #     # Tools
    #     tb = ttk.Frame(lt_fr); tb.pack(fill='x', pady=2)
    #     ttk.Button(tb, text="B", width=3, command=self.toggle_bold).pack(side='left')
    #     ttk.Button(tb, text="Color", width=5, command=self.pick_text_color).pack(side='left', padx=2)
        
    #     self.terms_txt = tk.Text(lt_fr, height=8, width=30, wrap='word', font=("Segoe UI", 9))
    #     self.terms_txt.pack(fill='both', expand=True, padx=5, pady=5)
    #     self.terms_txt.insert("1.0", "1. Payment: 100% Advance\n2. Validity: 30 Days\n3. Warranty: As per Policy")

    #     # --- COL 3: FOOTER & SIGNATURE ---
    #     col3 = ttk.Frame(bot_box)
    #     col3.grid(row=0, column=2, sticky="nsew", padx=10)
        
    #     sig_fr = ttk.LabelFrame(col3, text="Signatures")
    #     sig_fr.pack(fill='x')
        
    #     ttk.Label(sig_fr, text="Prepared By:").pack(anchor='w')
    #     ttk.Entry(sig_fr, textvariable=self.made_by_var).pack(fill='x', pady=2)
        
    #     ttk.Label(sig_fr, text="Approved By:").pack(anchor='w')
    #     ttk.Entry(sig_fr, textvariable=self.approved_by_var).pack(fill='x', pady=2)
    #     # 2. FOOTER SETTINGS (Ye Naya Code Hai)
    #     foot_fr = ttk.LabelFrame(col3, text="Page Footer / Sticker")
    #     foot_fr.pack(fill='x', pady=10)
        
    #     f_row = ttk.Frame(foot_fr)
    #     f_row.pack(fill='x', pady=2)
        
    #     # Buttons: Add Logo & Reset
    #     ttk.Button(f_row, text="+ Footer Img", command=lambda: self.load_logo("footer"), width=12).pack(side='left')
    #     ttk.Button(f_row, text="Reset", command=lambda: self.clear_logo("footer"), width=6).pack(side='left', padx=5)
        
    #     # Logo Status Label
    #     self.footer_logo_lbl = ttk.Label(f_row, text="(No Logo)", foreground="grey", font=("Arial", 8))
    #     self.footer_logo_lbl.pack(side='left', padx=5)

    #     # Footer Text
    #     ttk.Entry(foot_fr, textvariable=self.footer_text_var).pack(fill='x', pady=2, padx=2)
        
    #     # Options
    #     f_opt = ttk.Frame(foot_fr)
    #     f_opt.pack(fill='x')
    #     ttk.Checkbutton(f_opt, text="Full Width", variable=self.footer_full_width_var).pack(side='left')
    #     ttk.Checkbutton(f_opt, text="Pin Bottom", variable=self.footer_pin_to_bottom_var).pack(side='left', padx=5)
    #     # Save clear_inputsFVButton
          
    #     ttk.Button(col3, text="👁 PREVIEW & SAVE", style="success.TButton", command=self.on_preview_click).pack(fill='x', pady=20, ipady=10)
    def pick_header_color(self):
        """User se color select karwa kar poore header par apply karta hai"""
        color = colorchooser.askcolor(title="Choose Header BG Color", initialcolor=self.header_color)[1]
        if color:
            self.header_color = color
            # Tamam header rows ka background update karein
            for row in self.header_rows:
                row['bg'] = color
            # UI ko refresh karein taake naya color nazar aaye
            self.refresh_custom_header_ui()
    def _build_bottom_section(self, parent):
        # Main Container (Thora gap diya taake saaf lage)
        bot_box = ttk.Frame(parent)
        bot_box.pack(fill='both', expand=True, pady=20)
        
        # 3 Columns Layout
        bot_box.columnconfigure(0, weight=1)
        bot_box.columnconfigure(1, weight=1)
        bot_box.columnconfigure(2, weight=1)

        col1 = ttk.Labelframe(bot_box, text=" Financial Settings ", bootstyle="info", padding=15)
        col1.grid(row=0, column=0, sticky="nsew", padx=10)
        
        # Currency & Tax Rate
        r1 = ttk.Frame(col1)
        r1.pack(fill='x', pady=5)
        ttk.Label(r1, text="Currency:", foreground="white").pack(side='left')
        ttk.Combobox(r1, textvariable=self.currency_var, values=["PKR", "USD", "EUR", "GBP"], width=5, bootstyle="success").pack(side='left', padx=5)
        self.currency_var.trace('w', self.update_currency_symbol)
        
        ttk.Label(r1, text="Global Tax %:", foreground="white").pack(side='left', padx=(10, 5))
        gst_e = ttk.Entry(r1, textvariable=self.gst_rate_var, width=6, bootstyle="warning")
        gst_e.pack(side='left')
        gst_e.bind('<FocusOut>', lambda e: self.recalc_all())

        # --- NEW: EDITABLE FINANCIAL SUMMARY ---
        ttk.Separator(col1, bootstyle="secondary").pack(fill='x', pady=10)
        
        # Helper for Summary Rows
        def add_sum_row(parent, label, val_var, chk_var, is_bold=False):
            f = ttk.Frame(parent)
            f.pack(fill='x', pady=2)
            
            # Checkbox (Print?)
            ttk.Checkbutton(f, variable=chk_var, bootstyle="round-toggle").pack(side='left', padx=(0,5))
            
            # Label
            font = ("Segoe UI", 10, "bold") if is_bold else ("Segoe UI", 10)
            ttk.Label(f, text=label, width=12, font=font).pack(side='left')
            
            # Entry (Editable)
            e = ttk.Entry(f, textvariable=val_var, justify='right', font=font, width=15)
            e.pack(side='right', fill='x', expand=True)
            return e

        add_sum_row(col1, "Sub Total:", self.subtotal_var, self.print_subtotal_var)
        add_sum_row(col1, "Total Tax:", self.tax_total_var, self.print_tax_var)
        grand_ent = add_sum_row(col1, "Grand Total:", self.grand_total_var, self.print_grand_total_var, is_bold=True)
        # grand_ent.configure(bootstyle="success")

        # Extra Fields Container
        ttk.Separator(col1, bootstyle="secondary").pack(fill='x', pady=10)
        ttk.Button(col1, text="+ Add Extra Field", bootstyle="secondary-outline", command=self.add_extra_field).pack(anchor='w', fill='x')
        self.extra_cont = ttk.Frame(col1)
        self.extra_cont.pack(fill='x', pady=5)

        # --- COL 2: TERMS & CONDITIONS ---
        col2 = ttk.Labelframe(bot_box, text=" Terms & Conditions ", bootstyle="info", padding=15)
        col2.grid(row=0, column=1, sticky="nsew", padx=10)
        
        # Toolbar
        tb = ttk.Frame(col2)
        tb.pack(fill='x', pady=(0, 5))
        ttk.Button(tb, text="B", width=3, bootstyle="secondary-outline", command=self.toggle_bold).pack(side='left')
        ttk.Button(tb, text="Color", width=5, bootstyle="secondary-outline", command=self.pick_text_color).pack(side='left', padx=2)
        
        # Text Box (White bg taake text parhna asaan ho)
        self.terms_txt = tk.Text(col2, height=8, width=30, wrap='word', font=("Segoe UI", 9), bg="#ecf0f1", fg="black", relief="flat")
        self.terms_txt.pack(fill='both', expand=True)
        self.terms_txt.insert("1.0", "1. Payment: 100% Advance\n2. Validity: 30 Days\n3. Warranty: As per Policy")

        # --- COL 3: SIGNATURES & FOOTER ---
        col3 = ttk.Labelframe(bot_box, text=" Finalization ", bootstyle="info", padding=15)
        col3.grid(row=0, column=2, sticky="nsew", padx=10)
        
        # Signatures
        sig_fr = ttk.Frame(col3)
        sig_fr.pack(fill='x', pady=5)
        
        # Left Side: Prepared By
        prep_fr = ttk.Frame(sig_fr)
        prep_fr.pack(side='left', expand=True, fill='x')
        ttk.Label(prep_fr, text="Prepared By:", font=("Segoe UI", 9, "bold"), foreground="#3498db").pack(anchor='w')
        ttk.Entry(prep_fr, textvariable=self.made_by_var, bootstyle="primary").pack(fill='x', padx=(0, 5))
        
        # Right Side: Approved By
        app_fr = ttk.Frame(sig_fr)
        app_fr.pack(side='left', expand=True, fill='x')
        ttk.Label(app_fr, text="Approved By:", font=("Segoe UI", 9, "bold"), foreground="#3498db").pack(anchor='w')
        ttk.Entry(app_fr, textvariable=self.approved_by_var, bootstyle="primary").pack(fill='x', padx=(9, 0))

        # FOOTER SETTINGS (Card Style)
        # FOOTER SETTINGS (Updated with Alignment and Size)
        foot_fr = ttk.Labelframe(col3, text=" Page Footer / Sticker ", bootstyle="secondary", padding=10)
        foot_fr.pack(fill='x', pady=5)
        
        f_row1 = ttk.Frame(foot_fr)
        f_row1.pack(fill='x', pady=2)
        
        ttk.Button(f_row1, text="+ Footer Img", bootstyle="success-link", command=lambda: self.load_logo("footer"), width=12).pack(side='left')
        ttk.Button(f_row1, text="Reset", bootstyle="danger-link", command=lambda: self.clear_logo("footer"), width=6).pack(side='left', padx=5)
        self.footer_logo_lbl = ttk.Label(f_row1, text="(None)", foreground="grey", font=("Arial", 8))
        self.footer_logo_lbl.pack(side='left', padx=5)

        # --- NAYA SECTION: Alignment aur Size Controls ---
        f_row2 = ttk.Frame(foot_fr)
        f_row2.pack(fill='x', pady=5)
        
        ttk.Label(f_row2, text="Align:", font=("Arial", 8)).pack(side='left')
        # Alignment Dropdown
        ttk.Combobox(f_row2, textvariable=self.footer_align_var, values=["Left", "Center", "Right"], width=7, state="readonly").pack(side='left', padx=5)
        
        ttk.Label(f_row2, text="Size:", font=("Arial", 8)).pack(side='left', padx=(10, 0))
        # Size Slider (f_logo_size_var ko control karega)
        ttk.Scale(f_row2, variable=self.f_logo_size_var, from_=0.5, to=7.5, bootstyle="success", length=80).pack(side='left', padx=5)

        # Footer Text
        ttk.Entry(foot_fr, textvariable=self.footer_text_var, font=("Arial", 9)).pack(fill='x', pady=5)
        
        
        
        # Footer Options (Toggles)
        f_opt = ttk.Frame(foot_fr)
        f_opt.pack(fill='x')
        ttk.Checkbutton(f_opt, text="Full Width", variable=self.footer_full_width_var, bootstyle="round-toggle").pack(side='left')
        ttk.Checkbutton(f_opt, text="Pin Bottom", variable=self.footer_pin_to_bottom_var, bootstyle="round-toggle").pack(side='right')
        # BIG SAVE BUTTON (Green Filled)
        # messagebox.showinfo("Preview", "Disclaimer: Please verify your calculations. If you find any ambiguity, double tap to edit the cell.")  
        if self.on_preview_click:
           messagebox.showinfo("Preview", "Disclaimer: Please verify your calculations. If you find any ambiguity, double tap to edit the cell.")                
        ttk.Button(col3, text="👁 PREVIEW & SAVE", bootstyle="success", command=self.on_preview_click).pack(fill='x', pady=(15, 0), ipady=8)
    # =========================================================================
    # HEADER MANAGEMENT
    # =========================================================================
    def _init_standard_header_rows(self):
        defaults = [
            ("Customer Name:", self.client_name_var, "Date:", self.doc_date_var),
            ("Address:", self.client_addr_var, "Validity:", self.doc_validity_var),
            ("Contact Person:", self.client_contact_var, "Revised On:", self.revised_date_var),
            ("Designation:", self.client_designation_var, "email :", self.client_email_var),
            ("Contact No:", self.client_phone_var, "RFQ No:", self.rfq_no_var),
            (self.header_l6_label_var, self.header_l6_val_var, self.header_r6_label_var, self.header_r6_val_var),
        ]
        
        for i, (l1, v1, l2, v2) in enumerate(defaults):
            row_cfg = {
                'type': 'standard',
                'bg': "#ffffff", 
                'fg': "#000000",
                'l_val': v1,
                'r_val': v2
            }
            if i == 5: # Row 6 is dynamic labels
                row_cfg['l_label_var'] = l1
                row_cfg['r_label_var'] = l2
            else:
                row_cfg['l_label'] = l1
                row_cfg['r_label'] = l2
                
            self.header_rows.append(row_cfg)

    def refresh_custom_header_ui(self):
        for w in self.header_grid_fr.winfo_children(): w.destroy()
        
        # Grid for fields
        g_fr = ttk.Frame(self.header_grid_fr)
        g_fr.pack(fill='x', padx=5, pady=5)
        g_fr.columnconfigure(1, weight=1)
        g_fr.columnconfigure(3, weight=1)

        def add_grid_field(parent, label, var, row, col):
            ttk.Label(parent, text=label, font=("Segoe UI", 9, "bold")).grid(row=row, column=col*2, sticky='e', padx=5, pady=2)
            ttk.Entry(parent, textvariable=var).grid(row=row, column=col*2+1, sticky='ew', padx=5, pady=2)

        # Custom helper for row 6 (editable labels)
        def add_editable_grid_field(parent, label_var, val_var, row, col):
            ttk.Entry(parent, textvariable=label_var, font=("Segoe UI", 9, "bold"), width=12).grid(row=row, column=col*2, sticky='e', padx=5, pady=2)
            ttk.Entry(parent, textvariable=val_var).grid(row=row, column=col*2+1, sticky='ew', padx=5, pady=2)

        # Row 0
        add_grid_field(g_fr, "Customer Name:", self.client_name_var, 0, 0)
        add_grid_field(g_fr, "Date:", self.doc_date_var, 0, 1)

        # Row 1
        add_grid_field(g_fr, "Address:", self.client_addr_var, 1, 0)
        add_grid_field(g_fr, "Validity:", self.doc_validity_var, 1, 1)

        # Row 2
        add_grid_field(g_fr, "Contact Person:", self.client_contact_var, 2, 0)
        add_grid_field(g_fr, "Revised On:", self.revised_date_var, 2, 1)

        # Row 3
        add_grid_field(g_fr, "Designation:", self.client_designation_var, 3, 0)
        add_grid_field(g_fr, "email:", self.client_email_var, 3, 1)
         
        # Row 4
        add_grid_field(g_fr, "Contact No:", self.client_phone_var, 4, 0)
        add_grid_field(g_fr, "RFQ No:", self.rfq_no_var, 4, 1)
        
        # Row 5 (Editable labels)
        add_editable_grid_field(g_fr, self.header_l6_label_var, self.header_l6_val_var, 5, 0)
        add_editable_grid_field(g_fr, self.header_r6_label_var, self.header_r6_val_var, 5, 1)
        
        # Add Widgets for Custom Rows
        custom_exists = False
        for i, row in enumerate(self.header_rows):
            if row['type'] == 'standard': continue
            custom_exists = True
            
            fr = ttk.Frame(self.header_grid_fr)
            fr.pack(fill='x', pady=2)
            
            if row['type'] == 'full':
                ttk.Label(fr, text="Full Row:", width=10, style="Bold.TLabel").pack(side='left')
                ttk.Entry(fr, textvariable=row['l_label_var'], width=15).pack(side='left', padx=2)
                ttk.Entry(fr, textvariable=row['l_val'], width=40).pack(side='left', padx=2, fill='x', expand=True)
                
            elif row['type'] == 'split':
                ttk.Label(fr, text="Split Row:", width=10, style="Bold.TLabel").pack(side='left')
                # Left
                ttk.Entry(fr, textvariable=row['l_label_var'], width=12).pack(side='left', padx=2)
                ttk.Entry(fr, textvariable=row['l_val'], width=20).pack(side='left', padx=2)
                ttk.Label(fr, text="|", font=('Arial', 12, 'bold')).pack(side='left', padx=5)
                # Right
                ttk.Entry(fr, textvariable=row['r_label_var'], width=12).pack(side='left', padx=2)
                ttk.Entry(fr, textvariable=row['r_val'], width=20).pack(side='left', padx=2)

        if not custom_exists:
            ttk.Label(self.header_grid_fr, text="No custom rows added. Use 'Manage Header Rows' to add.", foreground="grey").pack()

    def open_header_manager(self):
        mgr = tk.Toplevel(self.root)
        mgr.title("Header Row Manager")
        mgr.geometry("900x600")
        mgr.transient(self.root)
        mgr.grab_set()

        ttk.Label(mgr, text="Manage Header Layout & Colors", font=('Segoe UI', 14, 'bold')).pack(pady=10)
        ttk.Label(mgr, text="Customize existing rows or add new Split (2-col) / Full (1-col) rows.", font=('Segoe UI', 10)).pack(pady=(0,10))

        list_fr = ttk.Frame(mgr)
        list_fr.pack(fill='both', expand=True, padx=20, pady=5)
        
        canvas = tk.Canvas(list_fr)
        sb = ttk.Scrollbar(list_fr, orient="vertical", command=canvas.yview)
        t_frame = ttk.Frame(canvas)
        
        t_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.create_window((0,0), window=t_frame, anchor="nw")
        canvas.configure(yscrollcommand=sb.set)
        
        canvas.pack(side="left", fill="both", expand=True)
        sb.pack(side="right", fill="y")
        
        self.mgr_canvas = canvas
        self.mgr_frame = t_frame
        
        btn_fr = ttk.Frame(mgr)
        btn_fr.pack(fill='x', padx=10, pady=10)
        ttk.Button(btn_fr, text="+ Add Split Row (Vendor | Client)", command=lambda: self.add_header_row("split")).pack(side='left', padx=5)
        ttk.Button(btn_fr, text="+ Add Full Row (Full Width)", command=lambda: self.add_header_row("full")).pack(side='left', padx=5)
        
        # CLOSE BUTTON REFRESHES MAIN UI
        def _close():
            self.refresh_custom_header_ui()
            mgr.destroy()

        ttk.Button(btn_fr, text="Done / Close", command=_close).pack(side='right')

        self.refresh_header_mgr()

    def refresh_header_mgr(self):
        for w in self.mgr_frame.winfo_children(): w.destroy()
        
        for idx, row in enumerate(self.header_rows):
            fr = ttk.LabelFrame(self.mgr_frame, text=f"Row {idx+1} ({row['type'].upper()})")
            fr.pack(fill='x', padx=5, pady=5)
            
            c_row = ttk.Frame(fr)
            c_row.pack(fill='x', pady=2)
            ttk.Button(c_row, text="BG Color", command=lambda r=row: self.set_header_color(r, 'bg')).pack(side='left')
            ttk.Label(c_row, text=f" {row['bg']} ", background=row['bg'], relief='solid').pack(side='left', padx=2)
            ttk.Button(c_row, text="Text Color", command=lambda r=row: self.set_header_color(r, 'fg')).pack(side='left', padx=(10,0))
            ttk.Label(c_row, text=f" {row['fg']} ", foreground=row['fg'], relief='solid').pack(side='left', padx=2)
            
            ttk.Button(c_row, text="🗑 Delete", command=lambda i=idx: self.delete_header_row(i)).pack(side='right', padx=10)

            cnt_row = ttk.Frame(fr)
            cnt_row.pack(fill='x', pady=5)
            
            if row['type'] == 'standard':
                if row.get('l_label'): 
                    ttk.Label(cnt_row, text=f"Left: {row['l_label']} [Standard Var]").pack(side='left')
                    ttk.Label(cnt_row, text=" | ").pack(side='left')
                    ttk.Label(cnt_row, text=f"Right: {row['r_label']} [Standard Var]").pack(side='left')
            
            elif row['type'] == 'split':
                ttk.Label(cnt_row, text="L-Lbl:").pack(side='left')
                ttk.Entry(cnt_row, textvariable=row['l_label_var'], width=10).pack(side='left', padx=2)
                ttk.Label(cnt_row, text="L-Val:").pack(side='left')
                ttk.Entry(cnt_row, textvariable=row['l_val'], width=15).pack(side='left', padx=2)
                ttk.Label(cnt_row, text="R-Lbl:").pack(side='left', padx=(10,0))
                ttk.Entry(cnt_row, textvariable=row['r_label_var'], width=10).pack(side='left', padx=2)
                ttk.Label(cnt_row, text="R-Val:").pack(side='left')
                ttk.Entry(cnt_row, textvariable=row['r_val'], width=15).pack(side='left', padx=2)
                
            elif row['type'] == 'full':
                ttk.Label(cnt_row, text="Label:").pack(side='left')
                ttk.Entry(cnt_row, textvariable=row['l_label_var'], width=15).pack(side='left', padx=2)
                ttk.Label(cnt_row, text="Value:").pack(side='left')
                ttk.Entry(cnt_row, textvariable=row['l_val'], width=40).pack(side='left', padx=2)

    def set_header_color(self, row, kind):
        c = colorchooser.askcolor(color=row[kind])[1]
        if c:
            row[kind] = c
            self.refresh_header_mgr()

    def add_header_row(self, kind):
        if kind == 'split':
            row = {
                'type': 'split', 'bg': '#ffffff', 'fg': '#000000',
                'l_label_var': tk.StringVar(value="Label 1"), 'l_val': tk.StringVar(),
                'r_label_var': tk.StringVar(value="Label 2"), 'r_val': tk.StringVar()
            }
        else:
            row = {
                'type': 'full', 'bg': '#ffffff', 'fg': '#000000',
                'l_label_var': tk.StringVar(value="Note:"), 'l_val': tk.StringVar(),
            }
        self.header_rows.append(row)
        self.refresh_header_mgr()

    def delete_header_row(self, idx):
        if 0 <= idx < len(self.header_rows):
            del self.header_rows[idx]
            self.refresh_header_mgr()

    # =========================================================================
    # COLUMN MANAGEMENT
    # =========================================================================
    def open_column_manager(self):
        # 1. Purani window clean karein (Overlap Fix)
        if hasattr(self, 'col_win') and self.col_win.winfo_exists():
            for child in self.col_win.winfo_children():
                child.destroy()
        else:
            self.col_win = tk.Toplevel(self.root)
        
        self.col_win.title("Advanced Column Manager")
        self.col_win.geometry("700x600")
        self.col_win.transient(self.root)
        self.col_win.grab_set()

        # 2. UI Elements
        ttk.Label(self.col_win, text="Drag & Drop Rows to Reorder", font=('Segoe UI', 14, 'bold')).pack(pady=10)
        
        main_fr = ttk.Frame(self.col_win, padding=10)
        main_fr.pack(fill='both', expand=True)
        
        cols = ("label", "width", "type", "print", "id")
        tv = ttk.Treeview(main_fr, columns=cols, show='headings', selectmode='browse', height=12)
        
        # Headings setup
        tv.heading("label", text="Header Label")
        tv.heading("width", text="Width")
        tv.heading("type", text="Type")
        tv.heading("print", text="Print?")
        tv.heading("id", text="ID")
        
        tv.column("label", width=200)
        tv.column("width", width=70, anchor='center')
        tv.column("type", width=100, anchor='center')
        tv.column("print", width=70, anchor='center')
        tv.column("id", width=100, anchor='center')
        
        tv.pack(side='left', fill='both', expand=True)

        # Enable Double Click to Edit
        def on_dbl_click(event):
            item_id = tv.identify_row(event.y)
            if not item_id: return
            idx = int(item_id)
            col_data = self.columns_config[idx]
            self._edit_col_dialog(col_data, lambda: self.open_column_manager(), self.col_win)

        tv.bind("<Double-1>", on_dbl_click)

        # --- THE DIRECT FIX: DATA YAHAN LOAD KAREIN ---
        # Kisi aur function par depend karne ke bajaye direct loop chalayein
        for idx, c in enumerate(self.columns_config):
            prt = "Yes" if c.get('printable', True) else "No"
            tv.insert("", "end", iid=str(idx), values=(c['label'], c['width'], c['type'], prt, c['id']))
        
        # --- ACTION BUTTONS ---
        btn_fr = ttk.Frame(self.col_win, padding=10)
        btn_fr.pack(fill='x')
        
        # Buttons ko functions ke saath sahi se link kiya hai
        ttk.Button(btn_fr, text="+ Add Column", bootstyle="success", command=lambda: self._edit_col_dialog(None, lambda: self.open_column_manager(), self.col_win)).pack(side='left', padx=5)
        ttk.Button(btn_fr, text="🗑 Delete Selected", bootstyle="danger", command=lambda: self.delete_column_from_tv(tv)).pack(side='left', padx=5)
        ttk.Button(btn_fr, text="✅ Save & Apply", bootstyle="primary", command=lambda: [self.refresh_items_ui_structure(), self.col_win.destroy()]).pack(side='right', padx=5)

    def delete_column_from_tv(self, tv):
        sel = tv.selection()
        if not sel: return
        idx = int(sel[0])
        if messagebox.askyesno("Confirm", f"Delete column '{self.columns_config[idx]['label']}'?", parent=self.col_win):
            del self.columns_config[idx]
            self.open_column_manager() # Window refresh karein
        
        def _refresh():
            for i in tv.get_children(): tv.delete(i)
            for idx, c in enumerate(self.columns_config):
                prt = "Yes" if c.get('printable', True) else "No"
                tv.insert("", "end", iid=str(idx), values=(c['label'], c['width'], c['type'], prt, c['id']))
        _refresh()
        
        # --- DRAG & DROP LOGIC ---
        def b_down(event):
            tv = event.widget
            if tv.identify_row(event.y) not in tv.get_children(): return
            tv.selection_set(tv.identify_row(event.y))
            
        def b_up(event):
            tv = event.widget
            moveto = tv.identify_row(event.y)
            dragged = tv.selection()
            if not dragged: return 
            
            dragged_idx = int(dragged[0])
            
            if moveto:
                 target_idx = int(moveto)
                 if target_idx != dragged_idx:
                     # Move element in config
                     item = self.columns_config.pop(dragged_idx)
                     self.columns_config.insert(target_idx, item)
                     _refresh()
                     tv.selection_set(str(target_idx))

        tv.bind("<ButtonPress-1>", b_down)
        tv.bind("<ButtonRelease-1>", b_up)

        # Tools
        btn_fr = ttk.Frame(win)
        btn_fr.pack(fill='x', pady=10)
        
        def _add(): self._edit_col_dialog(None, _refresh, win)
        def _edit():
            sel = tv.selection()
            if not sel: return
            idx = int(sel[0])
            self._edit_col_dialog(self.columns_config[idx], _refresh, win)
        def _delete():
            sel = tv.selection()
            if not sel: return
            idx = int(sel[0])
            c = self.columns_config[idx]
            if messagebox.askyesno("Confirm Delete", f"Delete column '{c['label']}'?"):
                del self.columns_config[idx]; _refresh()

        ttk.Button(btn_fr, text="Note: Drag rows to reorder", state="disabled").pack(side='left')
        ttk.Separator(btn_fr, orient='vertical').pack(side='left', padx=10, fill='y')
        ttk.Button(btn_fr, text="✏ Edit/Rename", command=_edit).pack(side='left', padx=2)
        ttk.Button(btn_fr, text="+ Add New", command=_add).pack(side='left', padx=2)
        ttk.Button(btn_fr, text="🗑 Delete", command=_delete).pack(side='left', padx=10)
        
        ttk.Button(win, text="Save Actions & Close", style="Action.TButton", command=lambda: [self.refresh_items_ui_structure(), win.destroy()]).pack(pady=10)

    def _edit_col_dialog(self, col_data, callback_fn, parent):
        d = tk.Toplevel(parent)
        d.title("Edit Column" if col_data else "Add Column")
        d.geometry("350x480")
        d.transient(parent); d.grab_set()
        
        is_new = col_data is None
        # Default 'printable' to True
        new_data = col_data.copy() if col_data else {'id': f"c_{len(self.columns_config)}_{int(datetime.datetime.now().timestamp())}", 'label': 'New Col', 'width': 100, 'type': 'text', 'printable': True}
        if 'printable' not in new_data: new_data['printable'] = True

        ttk.Label(d, text="Header Label:").pack(pady=(10,0))
        lbl_var = tk.StringVar(value=new_data['label'])
        ttk.Entry(d, textvariable=lbl_var).pack()
        
        ttk.Label(d, text="Width (px):").pack(pady=(5,0))
        w_var = tk.IntVar(value=new_data['width'])
        ttk.Entry(d, textvariable=w_var).pack()
        
        ttk.Label(d, text="Column Type:").pack(pady=(5,0))
        t_var = tk.StringVar(value=new_data['type'])
        cb = ttk.Combobox(d, textvariable=t_var, values=["text", "number", "calc", "global_pct", "auto"], state="readonly")
        cb.pack()
        
        # Printable Checkbox
        print_var = tk.BooleanVar(value=new_data['printable'])
        ttk.Checkbutton(d, text="Include in PDF/Word Output?", variable=print_var).pack(pady=10)
        
        ttk.Label(d, text="Note: 'calc' = Row Total\n'global_pct' = % of Amount", justify="center").pack(pady=10, padx=5)

        def _save():
            new_data['label'] = lbl_var.get()
            new_data['width'] = w_var.get()
            new_data['type'] = t_var.get()
            new_data['printable'] = print_var.get()
            
            if is_new: self.columns_config.append(new_data)
            else: 
                # Update in place
                col_data.update(new_data)
            callback_fn()
            d.destroy()
            
        ttk.Button(d, text="Save / Add Column", command=_save, style="Action.TButton").pack(pady=20, fill='x', padx=50)

    # quotation.py mein is function ko is se replace karein
# Yeh lines class ke andar (indented) honi chahiye
    def refresh_items_ui_structure(self):
        """Fixes the UI duplication bug by properly clearing and rebuilding the items section."""
        if hasattr(self, 'items_section_frame') and self.items_section_frame:
            # 1. Destroy all current children (Toolbar + Content)
            for widget in self.items_section_frame.winfo_children():
                widget.destroy()
            
            # 2. Re-build Toolbar
            tool_fr = ttk.Frame(self.items_section_frame)
            tool_fr.pack(fill='x', pady=5)
            ttk.Button(tool_fr, text="📊 Manage Columns", bootstyle="info-outline", command=self.open_column_manager).pack(side='left', padx=5)
            self._build_visibility_toolbar(tool_fr)

            # 3. Re-build Content (Inputs + Tree)
            self._build_items_section_content(self.items_section_frame)
            
            # 4. Final Refresh
            self.refresh_tree()
            self.recalc_all()

    def on_tree_double_click(self, event):
        """Opens a popup dialog to edit the cell content comfortably."""
        region = self.tree.identify_region(event.x, event.y)
        if region != "cell": return

        column_id = self.tree.identify_column(event.x) # e.g. "#1"
        item_id = self.tree.identify_row(event.y)
        
        # Identify Column & Row
        col_index = int(column_id[1:]) - 1
        real_col_id = self.tree["columns"][col_index]
        
        # Get Current Value
        current_val = self.tree.item(item_id)['values'][col_index]
        
        # Create Popup
        edit_win = tk.Toplevel(self.root)
        edit_win.title(f"Edit {real_col_id.title()}")
        edit_win.geometry("700x600")
        edit_win.transient(self.root)
        edit_win.grab_set()
        
        tk.Label(edit_win, text=f"Edit {real_col_id.title()}:", font=("Segoe UI", 10, "bold")).pack(pady=10)
        
        # Input Widget (Text area for description, Entry for others)
        if real_col_id in ['desc', 'description']:
            txt_input = tk.Text(edit_win, height=8, width=40, font=("Segoe UI", 10))
            txt_input.pack(padx=20, pady=5)
            txt_input.insert("1.0", str(current_val))
            input_widget = txt_input
            def get_val(): return txt_input.get("1.0", "end-1c")
        else:
            ent_input = ttk.Entry(edit_win, font=("Segoe UI", 10))
            ent_input.pack(padx=20, pady=5, fill='x')
            ent_input.insert(0, str(current_val))
            input_widget = ent_input
            def get_val(): return ent_input.get()
            
        input_widget.focus_set()

        def save_and_close(event=None):
            new_val = get_val()
            
            # 1. Update Backend Data (CRITICAL for PDF)
            row_idx = self.tree.index(item_id)
            if row_idx < len(self.items_data):
                # Number handling
                try: 
                    clean_val = float(new_val.replace(',', ''))
                    if clean_val.is_integer(): clean_val = int(clean_val)
                except: 
                    clean_val = new_val

                self.items_data[row_idx][real_col_id] = clean_val
                
                # Mark as Overridden (to prevent auto-recalc reverting it)
                if real_col_id in ['amount', 'gst', 'total', 'sno']:
                    overrides = self.items_data[row_idx].get('_overrides', [])
                    if real_col_id not in overrides:
                        overrides.append(real_col_id)
                    self.items_data[row_idx]['_overrides'] = overrides
            
            # 2. Refresh UI
            self.refresh_tree()
            self.recalc_all()
            edit_win.destroy()

        # Buttons
        btn_fr = ttk.Frame(edit_win)
        btn_fr.pack(pady=20)
        ttk.Button(btn_fr, text="✅ Confirm & Save", bootstyle="success", command=save_and_close).pack(side='left', padx=10)
        ttk.Button(btn_fr, text="❌ Cancel", bootstyle="secondary", command=edit_win.destroy).pack(side='left', padx=10)
        
        edit_win.bind("<Return>", lambda e: save_and_close() if real_col_id not in ['desc', 'description'] else None)
        edit_win.bind("<Escape>", lambda e: edit_win.destroy())

    def _on_header_right_click(self, event):
        region = self.tree.identify_region(event.x, event.y)
        if region != "heading": return
        
        col_id = self.tree.identify_column(event.x)
        col_idx = int(col_id.replace('#', '')) - 1
        if col_idx < 0 or col_idx >= len(self.columns_config): return
        
        c = self.columns_config[col_idx]
        
        # Context Menu
        menu = tk.Menu(self.root, tearoff=0)
        
        # Dynamic label based on state
        is_printable = c.get('printable', True)
        label = "Include in PDF/DOCX: YES" if is_printable else "Include in PDF/DOCX: NO"
        
        def toggle_print():
            c['printable'] = not is_printable
            self.refresh_items_ui_structure()
            # Maybe show a toast/message?
            state = "Enabled" if c['printable'] else "Disabled"
            messagebox.showinfo("Column Visibility", f"Printing {state} for column: {c['label']}")

        menu.add_command(label=label, command=toggle_print)
        menu.add_separator()
        menu.add_command(label="Manage Columns...", command=self.open_column_manager)
        
        menu.post(event.x_root, event.y_root)

    def _on_header_click(self, event):
        region = self.tree.identify_region(event.x, event.y)
        if region != "heading": return
        
        col_id = self.tree.identify_column(event.x)
        col_idx = int(col_id.replace('#', '')) - 1
        if col_idx < 0 or col_idx >= len(self.columns_config): return
        
        c = self.columns_config[col_idx]
        self._toggle_column_visibility(c)

    def _toggle_column_visibility(self, col_cfg):
        # Toggle Print Status
        col_cfg['printable'] = not col_cfg.get('printable', True)
        # Update UI components
        self._refresh_visibility_toolbar()
        sym = " ✔️" if col_cfg['printable'] else " ❌"
        self.tree.heading(col_cfg['id'], text=col_cfg['label'] + sym)

    def _build_visibility_toolbar(self, parent):
        self.viz_toolbar_frame = ttk.Frame(parent)
        self.viz_toolbar_frame.pack(fill='x', pady=(0, 5))
        self._refresh_visibility_toolbar()

    def _refresh_visibility_toolbar(self):
        if not hasattr(self, 'viz_toolbar_frame'): return
        for w in self.viz_toolbar_frame.winfo_children(): w.destroy()
        
        ttk.Label(self.viz_toolbar_frame, text="Include Columns in Output:", font=("Trebuchet MS", 10, "bold")).pack(side='left', padx=(5,10))
        
        for c in self.columns_config:
            is_prt = c.get('printable', True)
            btn_text = f"{'✔️' if is_prt else '❌'} {c['label']}"
            bg_color = "#e6fffa" if is_prt else "#fff5f5" # light teal vs light red
            fg_color = "#2c7a7b" if is_prt else "#c53030"
            
            # Using tk.Button for better color control than ttk
            btn = tk.Button(self.viz_toolbar_frame, text=btn_text, 
                            bg=bg_color, fg=fg_color,
                            font=("Trebuchet MS", 9), relief='flat',
                            command=lambda conf=c: self._toggle_column_visibility(conf))
            btn.pack(side='left', padx=2)
    


    # =========================================================================
    # ITEM MANAGEMENT & GRID
    # =========================================================================
    
    def _build_items_section_content(self, parent):
        # 1. GLOBAL RATES INPUTS
        gl_cols = [c for c in self.columns_config if c['type'] == 'global_pct']
        if gl_cols:
            gl_fr = ttk.LabelFrame(parent, text="Global Calculations")
            gl_fr.pack(fill='x', pady=(0, 10))
            for c in gl_cols:
                if c['id'] not in self.global_rates: 
                    self.global_rates[c['id']] = tk.DoubleVar(value=0.0)
                f = ttk.Frame(gl_fr)
                f.pack(side='left', padx=10)
                ttk.Label(f, text=c['label'] + " (%):").pack(side='left')
                ttk.Entry(f, textvariable=self.global_rates[c['id']], width=6).pack(side='left', padx=5)

        # 2. MAIN TOOLS LAYOUT (Left: Inputs | Right: Buttons)
        tools_layout = ttk.Frame(parent)
        tools_layout.pack(fill='x', pady=(10, 0))

        # --- LEFT SIDE: Add / Edit Item ---
        input_fr = ttk.LabelFrame(tools_layout, text="Add / Edit Item")
        input_fr.pack(side='left', fill='both', expand=True, padx=(0, 10))
        
        self.dynamic_vars = {}
        inp_cont = ttk.Frame(input_fr, padding=10)
        inp_cont.pack(fill='both', expand=True)
        
        current_row = ttk.Frame(inp_cont)
        current_row.pack(fill='x', pady=5)
        
        # --- FIXED: Label Logic aur New Col ka Masla ---
        for c in self.columns_config:
            if c['type'] in ['auto', 'calc', 'global_pct']: 
                continue 
            
            # Sirf in main columns ke labels nazar ayenge
            main_cols = ['desc', 'uom', 'qty', 'price']
            if c['id'] in main_cols:
                lbl_text = f"{c['label']}:"
                lbl = ttk.Label(current_row, text=lbl_text)
                lbl.pack(side='left', padx=(5,2))
            
            if c['id'] == 'desc':
                # Description Box bara aur wrap mode mein
                ent = tk.Text(current_row, width=45, height=4, font=("Segoe UI", 10), wrap='word')
                ent.pack(side='left', padx=(0, 10))
                self.dynamic_vars[c['id']] = ent 
            else:
                # Baqi inputs (labels ke baghair naye columns yahan ayenge)
                var = tk.StringVar()
                self.dynamic_vars[c['id']] = var
                width = 10 if c['type'] == 'number' else 15
                ent = ttk.Entry(current_row, textvariable=var, width=width)
                ent.pack(side='left', padx=(0, 5))

        # Action Buttons (Add, Clear, Bold)
        btn_row = ttk.Frame(input_fr)
        btn_row.pack(fill='x', pady=(0, 10), padx=10)
        ttk.Button(btn_row, textvariable=self.btn_add_text, command=self.add_or_update_item, bootstyle="success").pack(side='left', padx=5)
        ttk.Button(btn_row, text="Clear", command=self.clear_inputs, bootstyle="secondary").pack(side='left', padx=5)
        if 'desc' in self.dynamic_vars:
            ttk.Button(btn_row, text="Bold Desc", command=self.bold_desc, bootstyle="info-outline").pack(side='left', padx=5)

        # --- RIGHT SIDE: Row Controls (Alignment Fix) ---
        ctrl_fr = ttk.LabelFrame(tools_layout, text="Row Controls")
        ctrl_fr.pack(side='right', fill='y', padx=(0, 5))
        
        c_inner = ttk.Frame(ctrl_fr, padding=5)
        c_inner.pack(fill='both', expand=True)

        # Edit aur Color aik line mein
        r1 = ttk.Frame(c_inner); r1.pack(fill='x', pady=1)
        ttk.Button(r1, text="✎ Edit", command=self.load_row_for_edit, width=6, bootstyle="info").pack(side='left', padx=1)
        ttk.Button(r1, text="Color", command=self.set_row_color, width=6, bootstyle="secondary").pack(side='left', padx=1)

        # Delete aur Up/Down Arrows aik line mein
        r2 = ttk.Frame(c_inner); r2.pack(fill='x', pady=1)
        ttk.Button(r2, text="❌ Del", command=self.delete_row, width=6, bootstyle="danger").pack(side='left', padx=1)
        ttk.Button(r2, text="↩️ Undo", command=self.undo_delete, width=6, bootstyle="warning-outline").pack(side='left', padx=1)

        arrow_fr = ttk.Frame(c_inner); arrow_fr.pack(fill='x', pady=1)
        ttk.Button(arrow_fr, text="⬆", command=lambda: self.move_row(-1), width=4).pack(side='left', padx=1, expand=True, fill='x')
        ttk.Button(arrow_fr, text="⬇", command=lambda: self.move_row(1), width=4).pack(side='left', padx=1, expand=True, fill='x')

        # 3. TREEVIEW (Niche wali table)
        cols = [c['id'] for c in self.columns_config]
        self.tree = ttk.Treeview(parent, columns=cols, show='headings', selectmode='browse', height=10)
        for c in self.columns_config:
            sym = " ✔️" if c.get('printable', True) else " ❌"
            self.tree.heading(c['id'], text=c['label'] + sym)
            self.tree.column(c['id'], width=c['width'], anchor="center" if c['type']=='auto' else ("e" if c['type'] in ['number', 'calc', 'global_pct'] else "w"))

        vsb = ttk.Scrollbar(parent, orient="vertical", command=self.tree.yview)
        self.tree.configure(yscrollcommand=vsb.set)
        self.tree.pack(side='top', fill='both', expand=True)
        self.tree.bind("<Double-1>", self.on_tree_double_click)

    def add_or_update_item(self):
        if self.editing_index is not None: self._update_existing_item()
        else: self.add_item()

    # --- Std methods ---
    def bold_desc(self):
        v = self.dynamic_vars.get('desc')
        if not v: return
        if isinstance(v, tk.Text):
            try:
                # If there's selection, wrap it.
                v.insert("sel.first", "<b>")
                v.insert("sel.last", "</b>")
            except tk.TclError:
                # No selection, just append
                v.insert("end", " <b></b>")
        else:
            current = v.get()
            v.set(current + " <b></b>")

    def update_currency_symbol(self, *args):
        m = {"PKR": "Rs.", "INR": "₹", "USD": "$", "EUR": "€"}
        self.currency_symbol_var.set(m.get(self.currency_var.get(), ""))
        self.recalc_all()

    def pick_header_color(self):
        c = colorchooser.askcolor(color=self.header_color)[1]
        if c: self.header_color = c

    def load_logo(self, which):
        path = filedialog.askopenfilename(filetypes=[("Images", "*.png;*.jpg;*.jpeg")])
        if not path: return
        if which == "header": self.header_logo_path = path
        elif which == "middle": self.middle_logo_path = path
        elif which == "client": self.client_logo_path = path
        else: self.footer_logo_path = path
        
        img = Image.open(path); img.thumbnail((40, 40)); photo = ImageTk.PhotoImage(img)
        if which == "header": lbl = self.header_logo_lbl
        elif which == "middle": lbl = self.middle_logo_lbl
        elif which == "client": lbl = self.client_logo_lbl
        else: lbl = self.footer_logo_lbl
        
        lbl.config(image=photo, text=""); lbl.image = photo

    def clear_logo(self, which):
        if which == "header": 
            lbl = self.header_logo_lbl; self.header_logo_path = None
        elif which == "middle":
            lbl = self.middle_logo_lbl; self.middle_logo_path = None
        elif which == "client": 
            lbl = self.client_logo_lbl; self.client_logo_path = None
        else: 
            lbl = self.footer_logo_lbl; self.footer_logo_path = None
        
        lbl.config(image='', text="(No Logo)")

    def add_or_update_item(self):
        if self.editing_index is not None: self._update_existing_item()
        else: self.add_item()

    def _get_input_val(self, key):
        v = self.dynamic_vars[key]
        if isinstance(v, tk.Text): return v.get("1.0", "end-1c").strip()
        return v.get()

    def add_item(self):
        item = {}
        for col in self.columns_config:
            if col['type'] == 'number':
                try:
                    val = self._get_input_val(col['id'])
                    item[col['id']] = float(val.replace(',', '')) if val else 0.0
                except ValueError:
                    messagebox.showerror("Error", f"Invalid number for {col['label']}")
                    return
            elif col['type'] == 'text':
                item[col['id']] = self._get_input_val(col['id'])
            elif col['id'] == 'sno':
                item['sno'] = len(self.items_data) + 1
        
        self.items_data.append(item)
        self.refresh_tree()
        self.clear_inputs()
        self.recalc_all()

    def _update_existing_item(self):
        if self.editing_index is None: return
        item = self.items_data[self.editing_index]
        for col in self.columns_config:
            if col['type'] in ['text', 'number']:
                item[col['id']] = self._get_input_val(col['id'])
        self.editing_index = None
        self.btn_add_text.set("ADD ITEM")
        self.refresh_tree()
        self.clear_inputs()
        self.recalc_all()

    def load_row_for_edit(self):
        sel = self.tree.selection()
        if not sel: return
        idx = self.tree.index(sel[0])
        self.editing_index = idx
        item = self.items_data[idx]
        for k, var in self.dynamic_vars.items():
            if k in item: var.set(item[k])
        self.btn_add_text.set("UPDATE ITEM")

    def clear_inputs(self):
        for key, v in self.dynamic_vars.items():
            # Check karein ke ye Text Widget hai ya simple Variable
            if isinstance(v, tk.Text): 
                v.delete("1.0", "end") # Text box saaf karne ka tareeqa
            elif hasattr(v, 'set'):
                v.set("") # Simple input saaf karne ka tareeqa
            elif hasattr(v, 'delete'):
                v.delete(0, 'end') # Ehtiyat ke liye
                
        self.editing_index = None
        self.btn_add_text.set("ADD ITEM")
    def delete_row(self, *args):
        sel = self.tree.selection()
        if not sel: 
            messagebox.showwarning("Selection", "First Select it ")
            return
        
        # Confirmation zaroori hai
        if not messagebox.askyesno("Confirm", "Are You sure you want delete it", parent=self.root):
            return

        idx = self.tree.get_children().index(sel[0])

        # --- CRITICAL FIX: Backup for Undo ---
        # Data aur index ko class variable mein save karein
        self.last_deleted_item = {
            'data': self.items_data[idx].copy(),
            'index': idx,
            'color': self.row_colors.get(idx)
        }

        # Actual Deletion
        del self.items_data[idx]
        
        # Row colors shift logic
        new_colors = {}
        for k, v in self.row_colors.items():
            if k < idx: new_colors[k] = v
            elif k > idx: new_colors[k-1] = v
        self.row_colors = new_colors

        self.refresh_tree() #
        self.recalc_all()   #
        
        # Success Toast
        ToastNotification(title="Deleted", message="Undo (↩️) Press undo to return it back.", duration=3000, bootstyle="warning").show_toast()

    def undo_delete(self, *args):
        # 1. Check karein ke backup mojood hai ya nahi
        if not hasattr(self, 'last_deleted_item') or not self.last_deleted_item:
            messagebox.showinfo("No Undo", "No row!to undo", parent=self.root)
            return

        # 2. Backup se index aur data nikalna
        idx = self.last_deleted_item['index']
        data = self.last_deleted_item['data']
        color = self.last_deleted_item.get('color')

        # 3. Data list mein wapis insert karein
        self.items_data.insert(idx, data)
        
        # 4. Color restore logic
        if color:
            # Purane colors ko aage shift karein
            new_colors = {}
            for k, v in self.row_colors.items():
                if k < idx: new_colors[k] = v
                else: new_colors[k+1] = v
            new_colors[idx] = color
            self.row_colors = new_colors

        # 5. UI refresh aur totals update
        self.refresh_tree()
        self.recalc_all()
        
        # 6. Cleanup backup
        self.last_deleted_item = None
        messagebox.showinfo("Success", "Successfully get the deleted item!")
         
    # =========================================================================
    # UTILITY & HELPER METHODS
    # =========================================================================
    def bold_desc(self):
        sel = self.tree.selection()
        if not sel: return
        idx = self.tree.index(sel[0])
        color_key = int(idx)
        if self.row_colors.get(color_key) == "bold":
            self.row_colors[color_key] = "normal"
        else:
            self.row_colors[color_key] = "bold"
        self.refresh_tree()

    def update_currency_symbol(self, *args):
        c = self.currency_var.get()
        if c == "PKR": self.currency_symbol_var.set("Rs.")
        elif c == "USD": self.currency_symbol_var.set("$")
        elif c == "EURO": self.currency_symbol_var.set("€")
        self.recalc_all()

    def pick_header_color(self):
        color = colorchooser.askcolor(title="Choose Header BG Color")[1]
        if color:
            self.header_color = color
            for row in self.header_rows:
                row['bg'] = color
            self.refresh_custom_header_ui()

    def load_logo(self, which):
        path = filedialog.askopenfilename(filetypes=[("Image Files", "*.png *.jpg *.jpeg *.bmp")])
        if path:
            if which == "header": 
                self.header_logo_path = path
                self.header_logo_lbl.config(text=os.path.basename(path), foreground="green")
            elif which == "middle":
                self.middle_logo_path = path
                self.middle_logo_lbl.config(text=os.path.basename(path), foreground="green")
            elif which == "client":
                self.client_logo_path = path
            elif which == "footer":
                self.footer_logo_path = path

    def clear_logo(self, which):
        if which == "header":
            self.header_logo_path = None
            self.header_logo_lbl.config(text="(No Logo)", foreground="grey")
        elif which == "middle":
            self.middle_logo_path = None
            self.middle_logo_lbl.config(text="(No Logo)", foreground="grey")
        elif which == "footer":
            self.footer_logo_path = None
            self.footer_logo_lbl.config(text="(No Logo)", foreground="grey")

    def refresh_tree(self):
        for i in self.tree.get_children(): self.tree.delete(i)
        tr = self.gst_rate_var.get() / 100.0
        
        for idx, item in enumerate(self.items_data):
            # DYNAMIC CALCULATION ENGINE
            # 1. Base Amount
            # 1. Basic Math
            try: qty = float(item.get('qty', 0))
            except: qty = 0
            try: price = float(item.get('price', 0))
            except: price = 0
            
            overrides = item.get('_overrides', [])

            # Amount Calculation
            if 'amount' in overrides:
                try: amt = float(item.get('amount', 0))
                except: amt = 0.0
            else:
                amt = qty * price
                item['amount'] = amt
            
            # Store base amount for this row (used for taxes)
            # item['amount'] = amt # ensure it's set (Already done above or preserved)
            
            # 2. Taxation & Global Pct
            # Standard Tax (GST)
            if 'gst' in overrides:
                try: gst_val = float(item.get('gst', 0))
                except: gst_val = 0.0
            else:
                gst_val = amt * tr
                item['gst'] = gst_val
            
            # Global Percentage Columns
            current_extra_sum = 0
            for c in self.columns_config:
                if c['type'] == 'global_pct':
                    # Global percent columns are currently strictly calculated. 
                    # If we want them editable, we'd need similar override logic.
                    # For now, leaving as strictly calculated per common usage unless requested.
                    rate = self.global_rates.get(c['id'], tk.DoubleVar(value=0.0)).get()
                    val = amt * (rate / 100.0)
                    item[c['id']] = val # Store calculated value in item
                    current_extra_sum += val
                    
            # 3. Total
            if 'total' in overrides:
                 try: tot = float(item.get('total', 0))
                 except: tot = 0.0
            else:
                tot = amt + gst_val + current_extra_sum
                item['total'] = tot

            # S.No
            if 'sno' in overrides:
                # Keep existing sno
                pass
            else:
                item['sno'] = idx + 1
            
            # Build Row Values ordered by current config
            vals = []
            for c in self.columns_config:
                val = item.get(c['id'], "")
                
                # If it's a number/calc/global_pct, format it
                if c['type'] in ['number', 'calc', 'global_pct']:
                    try:
                        vals.append(f"{float(val):,.2f}")
                    except:
                         vals.append(str(val))
                else:
                    vals.append(str(val))
            
            iid = self.tree.insert("", "end", values=vals)
            if idx in self.row_colors:
                self.tree.item(iid, tags=(f"c_{idx}",)); self.tree.tag_configure(f"c_{idx}", background=self.row_colors[idx])

    def recalc_all(self):
        self.refresh_tree()
        
        # Calculate Totals
        sub_total = 0.0
        tax_total = 0.0
        
        for i in self.items_data:
            # Amount without Tax
            a = i.get('amount', 0)
            try: a = float(a)
            except: a = 0.0
            sub_total += a
            
            # Tax
            t = i.get('gst', 0)
            try: t = float(t)
            except: t = 0.0
            tax_total += t
            
            # Add Extra Calculates (Calculated Columns)
            # Logic: If 'total' is overridden, we trust it? 
            # Ideally: Grand Total = SubTotal + Tax + GlobalPCT?
            # But currently `total` row item includes everything.
            
        grand_total = sum(i.get('total', 0) for i in self.items_data)

        # Update Variables (Formatting .2f)
        self.subtotal_var.set(f"{sub_total:,.2f}")
        self.tax_total_var.set(f"{tax_total:,.2f}")
        self.grand_total_var.set(f"{grand_total:,.2f}")

        # Note: self.total_lbl Removed. Now using specific variables.

    def move_row(self, direction):
        sel = self.tree.selection()
        if not sel: return
        idx = self.tree.get_children().index(sel[0]); new_idx = idx + direction
        if 0 <= new_idx < len(self.items_data):
            self.items_data[idx], self.items_data[new_idx] = self.items_data[new_idx], self.items_data[idx]
            c1, c2 = self.row_colors.get(idx), self.row_colors.get(new_idx)
            if c1: self.row_colors[new_idx] = c1
            else: self.row_colors.pop(new_idx, None)
            if c2: self.row_colors[idx] = c2
            else: self.row_colors.pop(idx, None)
            self.refresh_tree(); self.tree.selection_set(self.tree.get_children()[new_idx])

    # def delete_row(self, *args):
    #     sel = self.tree.selection()
    #     if not sel: return
    #     idx = self.tree.get_children().index(sel[0])
    #     del self.items_data[idx]
    #     new_colors = {}
    #     for k, v in self.row_colors.items():
    #         if k < idx: new_colors[k] = v
    #         elif k > idx: new_colors[k-1] = v
    #     self.row_colors = new_colors; self.refresh_tree(); self.recalc_all()

    def set_row_color(self):
        sel = self.tree.selection()
        if not sel: return
        idx = self.tree.index(sel[0])
        c = colorchooser.askcolor()[1]
        if c: self.row_colors[idx] = c; self.refresh_tree()

    def add_extra_field(self):
        name = simpledialog.askstring("Field Name", "Enter label")
        if name:
            fr = ttk.Frame(self.extra_cont); fr.pack(fill='x', pady=1)
            ttk.Label(fr, text=f"{name}:").pack(side='left', padx=5)
            var = tk.StringVar(); ttk.Entry(fr, textvariable=var).pack(side='left', fill='x', expand=True)
            self.extra_fields.append((name, var))

    def toggle_bold(self):
        try:
            sel_start = self.terms_txt.index("sel.first")
            sel_end = self.terms_txt.index("sel.last")
            
            # Simple toggle: check first char
            all_tags = self.terms_txt.tag_names(sel_start)
            turning_on = "bt" not in all_tags
            
            if turning_on: self.terms_txt.tag_add("bt", sel_start, sel_end)
            else: self.terms_txt.tag_remove("bt", sel_start, sel_end)
            
            self._sync_tc_styles(sel_start, sel_end)
        except tk.TclError: pass

    def pick_text_color(self):
        c = colorchooser.askcolor()[1]
        if c: 
            try: 
                tag_name = f"ct_{c.replace('#', '')}"
                self.terms_txt.tag_add(tag_name, "sel.first", "sel.last")
                self.terms_txt.tag_config(tag_name, foreground=c)
                self._sync_tc_styles("sel.first", "sel.last")
            except: pass

    def change_font(self, event=None):
        try:
            sel_start = self.terms_txt.index("sel.first")
            sel_end = self.terms_txt.index("sel.last")
            f = self.font_family_var.get()
            
            for t in self.terms_txt.tag_names():
                if t.startswith("font_"): self.terms_txt.tag_remove(t, sel_start, sel_end)
            
            self.terms_txt.tag_add(f"font_{f}", sel_start, sel_end)
            self._sync_tc_styles(sel_start, sel_end)
        except tk.TclError: pass

    def change_size(self, event=None):
        try:
            sel_start = self.terms_txt.index("sel.first")
            sel_end = self.terms_txt.index("sel.last")
            s = self.font_size_var.get()
            
            for t in self.terms_txt.tag_names():
                if t.startswith("size_"): self.terms_txt.tag_remove(t, sel_start, sel_end)
            
            self.terms_txt.tag_add(f"size_{s}", sel_start, sel_end)
            self._sync_tc_styles(sel_start, sel_end)
        except tk.TclError: pass

    def _sync_tc_styles(self, start, end):
        
        txt = self.terms_txt
        start_idx = txt.index(start)
        end_idx = txt.index(end)
        
        # Get global defaults just in case
        def_f = self.font_family_var.get()
        try: def_s = int(self.font_size_var.get())
        except: def_s = 10

        curr = start_idx
        while txt.compare(curr, "<", end_idx):
            tags = txt.tag_names(curr)
            
            f = def_f
            s = def_s
            is_b = "bt" in tags
            
            for t in tags:
                if t.startswith("font_"): f = t.split("_")[1]
                elif t.startswith("size_"): s = int(t.split("_")[1])
            
            # Create a unique tag for this combination
            weight = "bold" if is_b else "normal"
            style_tag = f"style_{f}_{s}_{weight}"
            
            # Remove any other style_ tags at this position
            for t in tags:
                if t.startswith("style_"): txt.tag_remove(t, curr, f"{curr} + 1c")
            
            txt.tag_add(style_tag, curr, f"{curr} + 1c")
            txt.tag_config(style_tag, font=(f, s, weight))
            txt.tag_raise(style_tag)
            
            curr = txt.index(f"{curr} + 1c")

    def update_tc_ui_spacing(self):
        s = self.tc_spacing_var.get()
        # Tkinter Text uses spacing1, spacing2, spacing3 (top, line, bottom)
        # We'll use spacing3 for paragraph gap feedback
        self.terms_txt.config(spacing3=int(s))

    def _get_tagged_text(self):
        txt = self.terms_txt
        content = txt.get("1.0", "end-1c")
        
        result = []
        last_props = {"b": False, "f": None, "s": None, "c": None}
        
        # Helper to open/close tags based on prop changes (LIFO order)
        def sync_xml(new_props):
            # If bold status changes, or any font prop changes, we close everything and reopen
            # This is the simplest way to guarantee correct nesting in ReportLab XML.
            if last_props != new_props:
                # Close in reverse order: font then bold
                if last_props["f"] or last_props["s"] or last_props["c"]:
                    result.append("</font>")
                if last_props["b"]:
                    result.append("</b>")
                
                # Open in correct order: bold then font
                if new_props["b"]:
                    result.append("<b>")
                if new_props["f"] or new_props["s"] or new_props["c"]:
                    attrs = []
                    if new_props["f"]: attrs.append(f'face="{new_props["f"]}"')
                    if new_props["s"]: attrs.append(f'size="{new_props["s"]}"')
                    if new_props["c"]: attrs.append(f'color="{new_props["c"]}"')
                    result.append(f'<font {" ".join(attrs)}>')

        for i in range(len(content)):
            char = content[i]
            idx = f"1.0 + {i} chars"
            tags = txt.tag_names(idx)
            
            # Determine effective properties
            curr_props = {"b": "bt" in tags, "f": None, "s": None, "c": None}
            for t in tags:
                if t.startswith("font_"): curr_props["f"] = t.split("_")[1]
                elif t.startswith("size_"): curr_props["s"] = t.split("_")[1]
                elif t.startswith("ct_"): curr_props["c"] = "#" + t.split("_")[1]
                elif t.startswith("style_"):
                    # style_Font_Size_Weight
                    p = t.split("_")
                    if len(p) >= 4:
                        curr_props["f"] = p[1]
                        curr_props["s"] = p[2]
                        if p[3] == "bold": curr_props["b"] = True

            if char == "\n":
                sync_xml({"b": False, "f": None, "s": None, "c": None}) # Close all
                result.append("<br/>")
                last_props.update({"b": False, "f": None, "s": None, "c": None})
                # We'll re-open on next char
                continue

            sync_xml(curr_props)
            
            # Escape & add
            if char == "<": result.append("&lt;")
            elif char == ">": result.append("&gt;")
            elif char == "&": result.append("&amp;")
            else: result.append(char)
            
            last_props.update(curr_props)
            
        # Final cleanup
        sync_xml({"b": False, "f": None, "s": None, "c": None})
        return "".join(result)

    # =========================================================================
    # DOCUMENT GENERATION
    # =========================================================================
    def on_preview_click(self):
        # 1. Temporary File Create karein
        try:
            # Secure temp file banayi
            import tempfile
            fd, tmp_path = tempfile.mkstemp(suffix=".pdf")
            os.close(fd) # File handle release karein taake ReportLab likh sake
            
            # 2. PDF Generate karein
            self._generate_pdf(tmp_path)
            
            # 3. PDF Open karein (Windows vs Others)
            if os.path.exists(tmp_path):
                if os.name == 'nt':  # Agar Windows hai
                    os.startfile(tmp_path)
                else:  # Agar Mac/Linux hai
                    webbrowser.open("file://" + tmp_path)
            else:
                messagebox.showerror("Error", "Preview file could not be created.")
                return

        except Exception as e:
            messagebox.showerror("Preview Error", f"Could not generate preview.\nDetails: {str(e)}")
            return
            
        # 4. Confirmation Popup (Ab ye PDF khulne ke baad ayega)
        top = tk.Toplevel(self.root)
        top.title("Confirm Preview")
        top.geometry("400x300")
        top.transient(self.root)
        top.grab_set()
        
        ttk.Label(top, text="Check the opened PDF Layout.", font=('Segoe UI', 12, 'bold')).pack(pady=10)
        ttk.Label(top, text="If it looks good, choose a format to save:", foreground="grey").pack(pady=(0,10))
        
        btn_fr = ttk.Frame(top)
        btn_fr.pack(fill='both', expand=True, padx=20, pady=10)
        
        ttk.Button(btn_fr, text="✅ Approve & Save as PDF", command=lambda: self._finish_save(top, "pdf")).pack(fill='x', pady=5)
        ttk.Button(btn_fr, text="📄 Approve & Save as DOCX", command=lambda: self._finish_save(top, "docx")).pack(fill='x', pady=5)
        ttk.Button(btn_fr, text="📊 Approve & Save as Excel", command=lambda: self._finish_save(top, "xlsx")).pack(fill='x', pady=5)
        
        ttk.Separator(btn_fr, orient='horizontal').pack(fill='x', pady=10)
        ttk.Button(btn_fr, text="❌ Cancel / Edit More", command=top.destroy).pack(fill='x', pady=5)

    def _finish_save(self, dialog, fmt):
        dialog.destroy()
        if fmt == "pdf": self.save_pdf()
        elif fmt == "docx": self.save_docx()
        elif fmt == "excel": self.save_excel()

    # --- GENERATORS WITH DYNAMIC HEADER ---
    def _get_scaled_image(self, path, width_inch):
        if not path: return Spacer(1,1)
        try:
            img = ImageReader(path); iw, ih = img.getSize(); aspect = ih / float(iw)
            w = width_inch * inch; h = w * aspect; return RLImage(path, width=w, height=h)
        except: return Spacer(1,1)

    def _generate_qr_code(self, data, size_inch=0.6):
        """Generate a QR code image and return a ReportLab Image object."""
        if not qrcode: 
            print("QR Code library not available")
            return None
        try:
            qr = qrcode.QRCode(version=1, box_size=10, border=1)
            qr.add_data(data)
            qr.make(fit=True)
            img = qr.make_image(fill_color="black", back_color="white")
            tmp = tempfile.mktemp(suffix=".png")
            img.save(tmp)
            print(f"QR Code generated at: {tmp}")
            return RLImage(tmp, width=size_inch*inch, height=size_inch*inch)
        except Exception as e:
            print(f"QR Code generation error: {e}")
            return None

    def _generate_pdf(self, path):
        from reportlab.platypus import KeepTogether # Block protection ke liye

        # 1. SETUP
        MARGIN = 27.5 
        CW = 540 
        doc = SimpleDocTemplate(path, pagesize=A4, rightMargin=MARGIN, leftMargin=MARGIN, 
                                topMargin=MARGIN, bottomMargin=2.0*inch) # Margin thora barhaya
        elements = []
        styles = getSampleStyleSheet()
        norm_style = styles['Normal']
        
        # 2. HEADER
        styles = getSampleStyleSheet()
        norm_style = styles['Normal']
        
        v_img = self._get_scaled_image(self.header_logo_path, self.v_logo_size_var.get())
        m_img = self._get_scaled_image(self.middle_logo_path, self.m_logo_size_var.get())
        header_right = [
            Paragraph("Quotation", ParagraphStyle('T', parent=norm_style, fontName='Times-BoldItalic', fontSize=28, alignment=TA_RIGHT)),
            Spacer(1, 25),
            Paragraph(f"<b>No:</b> {self.quotation_no_var.get()}<br/><b>Date:</b> {self.doc_date_var.get()}", 
                      ParagraphStyle('N', parent=norm_style, alignment=TA_RIGHT, fontSize=10))
        ]
        t_header = Table([[v_img, m_img, header_right]], colWidths=[CW*0.3, CW*0.3, CW*0.4])
        t_header.setStyle(TableStyle([
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('ALIGN', (2,0), (2,0), 'RIGHT'),
            ('TOPPADDING', (2,0), (2,0), 45),
        ]))
        elements.append(t_header)
        elements.append(Spacer(1, 15))
        elements.append(Paragraph(f"<b>Vendor's Code: {self.vendor_code_var.get()}</b>", ParagraphStyle('C', parent=norm_style, alignment=TA_CENTER, fontSize=12)))
        elements.append(Spacer(1, 15))

        # 3. DYNAMIC TABLE
        print_cols = [c for c in self.columns_config if c.get('printable', True)]
        # Width logic (fixed for symmetry)
        # DYNAMIC COLUMN WIDTH LOGIC: PROPORTIONAL SCALING
        CW = 540 # Content Width
        
        # 1. Get Raw Widths from Config
        raw_widths = [float(c.get('width', 50)) for c in print_cols]
        total_raw = sum(raw_widths)
        
        # 2. Calculate Scale Factor
        if total_raw == 0: total_raw = 1 
        scale_factor = CW / total_raw
        
        # 3. Apply Scale Factor to create PDF Widths
        pdf_widths = [w * scale_factor for w in raw_widths]

        data = [[Paragraph(f"<b>{c['label']}</b>", ParagraphStyle('H', parent=norm_style, fontSize=9, alignment=TA_CENTER)) for c in print_cols]]
        for item in self.items_data:
            row = []
            for c in print_cols:
                val = item.get(c['id'], "")
                # Leading adjust kiya taake text wrap na ho
                p_s = ParagraphStyle('Cell', parent=norm_style, fontSize=8, leading=9, alignment=TA_RIGHT if c['type'] in ['number','calc','global_pct'] else TA_LEFT)
                
                if c['type'] in ['number', 'calc', 'global_pct']:
                    try: val = f"{float(val):,.2f}"
                    except: pass
                
                row.append(Paragraph(str(val), p_s))
            data.append(row)

        t_items = Table(data, colWidths=pdf_widths, repeatRows=1)
        t_items.setStyle(TableStyle([
            ('GRID', (0,0), (-1,-1), 0.5, colors.black), 
            ('BACKGROUND', (0,0), (-1,0), self.header_color),
            ('TOPPADDING', (0,0), (-1,-1), 4),
            ('BOTTOMPADDING', (0,0), (-1,-1), 4),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ]))
        elements.append(t_items)
        elements.append(Spacer(1, 10))

        # --- FINANCIAL SUMMARY TABLE (Conditional) ---
        summary_data = []
        curr = self.currency_symbol_var.get()
        
        # Sub Total
        if self.print_subtotal_var.get():
            summary_data.append([
                Paragraph("<b>Total Amount (Excl. Tax):</b>", ParagraphStyle('SL', parent=norm_style, alignment=TA_RIGHT)),
                Paragraph(f"<b>{curr} {self.subtotal_var.get()}</b>", ParagraphStyle('SV', parent=norm_style, alignment=TA_RIGHT))
            ])
            
        # Tax Total
        if self.print_tax_var.get():
            summary_data.append([
                Paragraph("<b>Total Sales Tax:</b>", ParagraphStyle('SL', parent=norm_style, alignment=TA_RIGHT)),
                Paragraph(f"<b>{curr} {self.tax_total_var.get()}</b>", ParagraphStyle('SV', parent=norm_style, alignment=TA_RIGHT))
            ])
            
        # Grand Total
        if self.print_grand_total_var.get():
            summary_data.append([
                Paragraph("<b>Grand Total:</b>", ParagraphStyle('SL', parent=norm_style, alignment=TA_RIGHT, fontSize=11)),
                Paragraph(f"<b>{curr} {self.grand_total_var.get()}</b>", ParagraphStyle('SV', parent=norm_style, alignment=TA_RIGHT, fontSize=11))
            ])
            
        if summary_data:
            # Align right side of page
            # Table width logic: let's say 40% of page width aligned right
            t_sum = Table(summary_data, colWidths=[CW*0.25, CW*0.15])
            t_sum.setStyle(TableStyle([
                ('LINEABOVE', (0,0), (-1,0), 1, colors.black),
                # ('GRID', (0,0), (-1,-1), 0.5, colors.grey), # Uncomment to debug grid
                ('ALIGN', (0,0), (-1,-1), 'RIGHT'),
                ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ]))
            # Use a wrapper table to push it to the right
            wrapper = Table([[ "", t_sum ]], colWidths=[CW*0.6, CW*0.4])
            wrapper.setStyle(TableStyle([('ALIGN', (1,0), (1,0), 'RIGHT')]))
            elements.append(wrapper)

        elements.append(Spacer(1, 25))

        # 4. TERMS & CONDITIONS (Restored)
        elements.append(Paragraph("<b>Terms & Conditions:</b>", ParagraphStyle('BT', parent=norm_style, fontSize=10)))
        # Use a spacer or keep it close
        elements.append(Paragraph(self._get_tagged_text(), ParagraphStyle('BC', parent=norm_style, fontSize=9)))
        elements.append(Spacer(1, 20))

        # 5. SIGNATURES BLOCK
        # Prepared / Approved By (Alignment Fix)
        footer_elements = []  # Start a new list for the keep-together block if needed, or just append to elements
        
        # We want "Prepared By" (Left), "System Note" (Center), "Approved By" (Right)
        # All on generally the same horizontal level, bottom-aligned.
        
        sig_style_left = ParagraphStyle('SigL', parent=norm_style, fontSize=10, alignment=TA_LEFT)
        sig_style_right = ParagraphStyle('SigR', parent=norm_style, fontSize=10, alignment=TA_RIGHT)
        sig_style_center = ParagraphStyle('SigC', parent=norm_style, fontSize=8, alignment=TA_CENTER, textColor=colors.red)
        
        # Structure:
        # Row 1: ___________       (Note)        ___________
        # Row 2: Prepared By   (System Gen...)   Approved By
        
        # Ideally, we can do it in one cell with <br/> or separate rows.
        # Let's use separate cells for cleanliness.
        
        is_quotation = self.__class__.__name__ == "QuotationApp"
        
        if is_quotation:
            # Note Text
            note_text = "Note: This is a system generated document so no need to sign."
            
            sig_data = [[
                Paragraph("_________________<br/><b>Prepared By</b>", sig_style_left),
                Paragraph(f"<br/>{note_text}", sig_style_center), 
                Paragraph("_________________<br/><b>Approved By</b>", sig_style_right)
            ]]
        else:
            sig_data = [[
                Paragraph("_________________<br/><b>Prepared By</b>", sig_style_left),
                "", 
                Paragraph("_________________<br/><b>Approved By</b>", sig_style_right)
            ]]
        
        # Adjust Column Widths to ensure "System Note" fits in center without squashing signatures
        t_sig = Table(sig_data, colWidths=[CW*0.3, CW*0.4, CW*0.3])
        t_sig.setStyle(TableStyle([
            ('VALIGN', (0,0), (-1,-1), 'BOTTOM'),
            ('ALIGN', (0,0), (0,0), 'LEFT'),
            ('ALIGN', (1,0), (1,0), 'CENTER'),
            ('ALIGN', (2,0), (2,0), 'RIGHT'),
        ]))
        
        # Wrap in KeepTogether to avoid breaking mid-signature
        elements.append(KeepTogether([t_sig]))
        # 5. CANVAS (Social Icons & QR)
        qr_img = self._generate_qr_code("https://www.orientmarketing.com.pk/", size_inch=0.5)
        
        def canvas_setup(canvas, doc):
            canvas.saveState()
            page_w, page_h = A4
            
            # --- FOOTER STICKER (IMAGE) ---
            if self.footer_logo_path:
                # Slider se size uthayega
                f_img = self._get_scaled_image(self.footer_logo_path, self.f_logo_size_var.get())
                if f_img:
                    w = f_img.drawWidth
                    # Dropdown se alignment uthayega
                    align = self.footer_align_var.get()
                    if align == "Center":
                        x_pos = (page_w - w) / 2
                    elif align == "Right":
                        x_pos = page_w - MARGIN - w
                    else: # Left
                        x_pos = MARGIN
                    # Image draw karein
                    f_img.drawOn(canvas, x_pos, 60) 
            
            # --- SOCIAL ICONS WITH LINKS ---
            try:
                social_x, social_y = MARGIN, 35
                icons = [
                    ('#0066cc', 'W', 'https://www.orientmarketing.com.pk/'), 
                    ('#FF0000', 'Y', 'https://www.youtube.com/@Antarc-Technologies'), 
                    ('#1877F2', 'f', 'https://www.facebook.com/orientmarketing.com.pk')
                ]
                for idx, (color, sym, url) in enumerate(icons):
                    x = social_x + idx * 22
                    canvas.setFillColor(color)
                    canvas.rect(x, social_y, 16, 16, fill=1, stroke=0)
                    canvas.setFillColor(colors.white)
                    canvas.setFont("Helvetica-Bold", 10)
                    canvas.drawString(x+4, social_y+4, sym)
                    canvas.linkURL(url, (x, social_y, x + 16, social_y + 16))
            except: pass
            
            if qr_img:
                qr_img.drawOn(canvas, page_w - MARGIN - 60, 30)
            canvas.restoreState()

        doc.build(elements, onFirstPage=canvas_setup, onLaterPages=canvas_setup)

    # --- ERROR HANDLING SAVERS ---
    def save_pdf(self):
        try:
            f = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF Documents", "*.pdf")], parent=self.root)
            if not f: return
            self._generate_pdf(f)
            if os.path.exists(f):
                pg_msg = f" ({self.last_pdf_pages} pages)" if self.last_pdf_pages > 1 else ""
                messagebox.showinfo("Success", f"PDF Successfully Saved{pg_msg} at:\n{f}")
            else:
                messagebox.showerror("Error", "File was not created on disk.")
        except Exception as e:
            messagebox.showerror("Save Error", f"Failed to save PDF:\n{str(e)}")
            print(f"PDF Error: {e}")

    def save_docx(self):
        if not docx: 
            messagebox.showerror("Error", "python-docx library not installed.")
            return
        try:
            f = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Documents", "*.docx")], parent=self.root)
            if not f: return
            
            doc = docx.Document()
            # Standardize Geometry
            for sec in doc.sections:
                sec.top_margin = Inches(0.5); sec.bottom_margin = Inches(0.5)
                sec.left_margin = Inches(0.5); sec.right_margin = Inches(0.5)
            CW_INCH = 7.5
            
            def get_align(val):
                if val == "Center": return WD_ALIGN_PARAGRAPH.CENTER
                if val == "Right": return WD_ALIGN_PARAGRAPH.RIGHT
                return WD_ALIGN_PARAGRAPH.LEFT

            def add_logo_to_cell(cell, path, size_val, align_val):
                if not path or not os.path.exists(path): return
                p = cell.paragraphs[0]
                p.alignment = get_align(align_val)
                run = p.add_run()
                run.add_picture(path, width=Inches(size_val))

            # 1. Header Structure: Logo (Left), Mid Logo (Center), Title/No (Right)
            t_header = doc.add_table(rows=1, cols=3)
            t_header.width = Inches(CW_INCH)
            t_header.autofit = False 
            t_header.columns[0].width = Inches(CW_INCH * 0.33)
            t_header.columns[1].width = Inches(CW_INCH * 0.34)
            t_header.columns[2].width = Inches(CW_INCH * 0.33)
            if WD_ALIGN_VERTICAL:
                for cell in t_header.rows[0].cells:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            # Left: Primary Logo
            add_logo_to_cell(t_header.cell(0, 0), self.header_logo_path, self.v_logo_size_var.get(), "Left")
            
            # Center: Middle Logo
            add_logo_to_cell(t_header.cell(0, 1), self.middle_logo_path, self.m_logo_size_var.get(), "Center")
            
            # Right: Quotation Title & No
            cell_right = t_header.cell(0, 2)
            p_title = cell_right.paragraphs[0]
            p_title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_title.paragraph_format.line_spacing = 1.0 
            run_title = p_title.add_run("Quotation")
            run_title.font.name = 'Segoe Script'
            run_title.font.size = Pt(28)
            run_title.bold = True
            
            p_no = cell_right.add_paragraph()
            p_no.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_no.paragraph_format.space_before = Pt(0)
            run_no = p_no.add_run(f"Qtn No: {self.quotation_no_var.get()}")
            run_no.font.bold = True
            run_no.font.size = Pt(11)

            doc.add_paragraph("") # Spacer

            # 2. Vendor Code (Centered)
            p_code = doc.add_paragraph()
            p_code.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_code = p_code.add_run(f"Vendor's Code: {self.vendor_code_var.get()}")
            run_code.font.bold = True
            run_code.font.size = Pt(12)
            
            doc.add_paragraph("") # Spacer

            
            # 3. DYNAMIC HEADER TABLE (Two-Column Layout)
            table = doc.add_table(rows=0, cols=4); table.style = 'Table Grid'
            
            # Data Rows (Smart Hiding)
            for row_cfg in self.header_rows:
                l_val = row_cfg['l_val'].get().strip() if 'l_val' in row_cfg else ""
                r_val = row_cfg['r_val'].get().strip() if 'r_val' in row_cfg else ""
                l_lbl = row_cfg['l_label_var'].get().strip() if 'l_label_var' in row_cfg else row_cfg.get('l_label', "").strip()
                r_lbl = row_cfg['r_label_var'].get().strip() if 'r_label_var' in row_cfg else row_cfg.get('r_label', "").strip()
                
                show_l = (l_lbl != "" if 'l_label_var' in row_cfg else l_val != "")
                show_r = (r_lbl != "" if 'r_label_var' in row_cfg else r_val != "")

                if not show_l and not show_r: continue
                
                r = table.add_row()
                if row_cfg['type'] == 'full':
                     r.cells[0].text = f"{l_lbl} {l_val}"
                     r.cells[0].merge(r.cells[3])
                else:
                    if show_l:
                        r.cells[0].text = l_lbl; r.cells[0].paragraphs[0].runs[0].bold=True
                        r.cells[1].text = l_val
                    if show_r:
                        r.cells[2].text = r_lbl; r.cells[2].paragraphs[0].runs[0].bold=True
                        r.cells[3].text = r_val
            
            doc.add_paragraph("")
            
            # --- DYNAMIC ITEMS DOCX ---
            # Filter printable
            print_cols = [c for c in self.columns_config if c.get('printable', True)]
            
            t_items = doc.add_table(rows=1, cols=len(print_cols)); t_items.style = 'Table Grid'
            hdr_cells = t_items.rows[0].cells
            for i, c in enumerate(print_cols):
                hdr_cells[i].text = c['label']
                hdr_cells[i].paragraphs[0].runs[0].bold = True
                
            # Set Header Repeat (Word Specific XML)
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            tr = t_items.rows[0]._tr
            trPr = tr.get_or_add_trPr()
            tblHeader = OxmlElement('w:tblHeader')
            tblHeader.set(qn('w:val'), "true")
            trPr.append(tblHeader)

            for item in self.items_data:
                row_cells = t_items.add_row().cells
                for i, c in enumerate(print_cols):
                    val = item.get(c['id'], "")
                    if c['type'] in ['number', 'calc', 'global_pct']:
                        try: row_cells[i].text = f"{float(val):,.2f}"
                        except: row_cells[i].text = str(val)
                    else:
                        row_cells[i].text = str(val)
            
            doc.add_paragraph(self.total_lbl.cget("text")).alignment = WD_ALIGN_PARAGRAPH.RIGHT
            doc.add_paragraph("")
            
            # --- Formatted Terms (Word Parser with custom spacing) ---
            doc.add_paragraph("Terms & Conditions:", style='Normal').runs[0].bold = True
            
            tagged_terms = self._get_tagged_text()
            tc_spc_pt = self.tc_spacing_var.get() * 12
            
            lines = tagged_terms.split('<br/>')
            for line in lines:
                if not line.strip(): continue
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(tc_spc_pt)
                
                parts = re.split(r'(</?[a-z][^>]*>)', line)
                bold_stack = [False]; color_stack = [None]; face_stack = [None]; size_stack = [None]
                
                for part in parts:
                    if not part: continue
                    if part == "<b>": bold_stack.append(True)
                    elif part == "</b>": 
                        if len(bold_stack) > 1: bold_stack.pop()
                    elif part.startswith("<font"):
                        color_stack.append(color_stack[-1]); face_stack.append(face_stack[-1]); size_stack.append(size_stack[-1])
                        c_m = re.search(r'color="([^"]*)"', part)
                        if c_m: color_stack[-1] = c_m.group(1)
                        f_m = re.search(r'face="([^"]*)"', part)
                        if f_m: face_stack[-1] = f_m.group(1)
                        s_m = re.search(r'size="([^"]*)"', part)
                        if s_m: size_stack[-1] = s_m.group(1)
                    elif part == "</font>":
                        if len(color_stack) > 1: color_stack.pop()
                        if len(face_stack) > 1: face_stack.pop()
                        if len(size_stack) > 1: size_stack.pop()
                    else:
                        txt_part = part.replace("&lt;", "<").replace("&gt;", ">").replace("&amp;", "&")
                        run = p.add_run(txt_part)
                        if bold_stack[-1]: run.bold = True
                        if color_stack[-1]:
                            c = color_stack[-1].lstrip('#')
                            try: run.font.color.rgb = RGBColor(int(c[0:2], 16), int(c[2:4], 16), int(c[4:6], 16))
                            except: pass
                        if face_stack[-1]: run.font.name = face_stack[-1]
                        if size_stack[-1]: 
                            try: run.font.size = Pt(int(size_stack[-1]))
                            except: pass

            if self.extra_fields:
                doc.add_paragraph("Additional Info:", style='Normal').runs[0].bold = True
                for n, v in self.extra_fields:
                    p = doc.add_paragraph() 
                    p.add_run(f"{n}: ").bold = True
                    p.add_run(v.get())
                    p.paragraph_format.space_after = Pt(6)

            # Signature Section
            doc.add_paragraph(""); doc.add_paragraph("")
            t_sig = doc.add_table(rows=1, cols=2)
            t_sig.width = Inches(CW_INCH); t_sig.autofit = False
            t_sig.columns[0].width = Inches(CW_INCH*0.5); t_sig.columns[1].width = Inches(CW_INCH*0.5)
            
            p_l = t_sig.cell(0, 0).paragraphs[0]
            p_l.add_run(f"Prepared By: {self.made_by_var.get()}").bold = True
             
            
            p_r = t_sig.cell(0, 1).paragraphs[0]
            p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_r.add_run(f"Approved By: {self.approved_by_var.get()}").bold = True
            
            fp = doc.sections[0].footer.paragraphs[0]
            fp.alignment = get_align(self.footer_align_var.get())
            if self.footer_logo_path:
                w_val = 7.5 if self.footer_full_width_var.get() else self.f_logo_size_var.get()
                fp.add_run().add_picture(self.footer_logo_path, width=Inches(w_val))
                fp.add_run("\n")
            fp.add_run(self.footer_text_var.get())
            doc.save(f)
            
            # Heuristic for multi-page in Word
            is_large = len(self.items_data) > 25
            pg_note = "\n(Multiple pages likely generated)" if is_large else ""
            messagebox.showinfo("Success", f"DOCX Saved successfully at:\n{f}{pg_note}")
            
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save DOCX: {e}")

    def save_excel(self):
        if not openpyxl: 
            messagebox.showerror("Error", "openpyxl library not installed.")
            return
        try:
            f = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel Files", "*.xlsx")], parent=self.root)
            if not f: return
            
            wb = openpyxl.Workbook(); ws = wb.active
            thin = Side(border_style="thin", color="000000"); border = Border(top=thin, left=thin, right=thin, bottom=thin)
            # 1. Header Layout
            ws.merge_cells("A1:B2"); ws['A1'] = "[LOGO 1]"; ws['A1'].alignment = Alignment(horizontal='center', vertical='center')
            ws.merge_cells("C1:E2"); ws['C1'] = "[MID LOGO]"; ws['C1'].alignment = Alignment(horizontal='center', vertical='center')
            
            ws.merge_cells("F1:H1"); ws['F1'] = "Quotation"; 
            ws['F1'].font = Font(size=24, bold=True, italic=True); ws['F1'].alignment = Alignment(horizontal='right')
            
            ws.merge_cells("F2:H2"); ws['F2'] = f"Quotation No: {self.quotation_no_var.get()}"
            ws['F2'].font = Font(bold=True); ws['F2'].alignment = Alignment(horizontal='right')

            ws.merge_cells("D3:E3"); ws['D3'] = f"Vendor's Code: {self.vendor_code_var.get()}"
            ws['D3'].font = Font(bold=True); ws['D3'].alignment = Alignment(horizontal='center')
            
            # Conditional Header in Excel (simplified)
            # Just print rows as they are.
            r=4
            for row_cfg in self.header_rows:
                l_val = row_cfg['l_val'].get().strip()
                r_val = row_cfg['r_val'].get().strip()
                
                if not l_val and not r_val: continue
                
                l_lbl = row_cfg['l_label_var'].get() if 'l_label_var' in row_cfg else row_cfg['l_label']
                r_lbl = row_cfg['r_label_var'].get() if 'r_label_var' in row_cfg else row_cfg['r_label']
                
                if l_val:
                    ws[f"A{r}"] = l_lbl; ws[f"B{r}"] = l_val
                if r_val:
                    ws[f"E{r}"] = r_lbl; ws[f"F{r}"] = r_val
                r+=1
            
            r += 2
            # Dynamic Columns
            # Filter printable
            print_cols = [c for c in self.columns_config if c.get('printable', True)]
            
            for i, c in enumerate(print_cols):
                cell = ws.cell(row=r, column=i+1, value=c['label'])
                cell.font = Font(bold=True); cell.border = border
            r += 1
            
            for item in self.items_data:
                for i, c in enumerate(print_cols):
                    val = item.get(c['id'], "")
                    try: val = float(val)
                    except: pass
                    ws.cell(row=r, column=i+1, value=val).border = border
                r += 1
            
            ws.cell(row=r+1, column=len(print_cols), value=self.total_lbl.cget("text")).font = Font(bold=True)
            wb.save(f)
            messagebox.showinfo("Success", f"Excel Saved: {f}")
        except Exception as e:
            messagebox.showerror("Save Error", f"Failed to save Excel:\n{str(e)}")


def start_app():
    root = tb.Window(themename="superhero")    
    root.withdraw() # Sab se pehle main app ko chhupa dein

    # --- FULL SCREEN SPLASH WINDOW ---
    splash = tk.Toplevel(root)
    splash.overrideredirect(True) 
    
    sw = root.winfo_screenwidth()
    sh = root.winfo_screenheight()
    splash.geometry(f"{sw}x{sh}+0+0") 
    splash.configure(bg="#121212") 
    splash.attributes('-topmost', True) # Screen ke upar rakhne ke liye

    # --- UI ELEMENTS SETUP ---
    # 1. Main Title
    title_lbl = tk.Label(splash, text="", font=("Segoe UI", 45, "bold"), fg="#00e676", bg="#121212")
    title_lbl.place(relx=0.5, rely=0.25, anchor='center')
    
    # 2. Subtitle (ODM-ONLINE)
    sub_lbl = tk.Label(splash, text="", font=("Segoe UI", 30, "bold"), fg="white", bg="#121212")
    sub_lbl.place(relx=0.5, rely=0.38, anchor='center')

    # 3. Slogan (Joining Writing)
    slogan_lbl = tk.Label(splash, text="", font=("Segoe Script", 18, "italic"), fg="#f1c40f", bg="#121212")
    slogan_lbl.place(relx=0.5, rely=0.48, anchor='center')

    # 4. Powered By
    pow_lbl = tk.Label(splash, text="", font=("Segoe UI", 18, "italic"), fg="#ecf0f1", bg="#121212")
    pow_lbl.place(relx=0.5, rely=0.58, anchor='center')

    # 5. Website Link
    web_lbl = tk.Label(splash, text="", font=("Consolas", 15), fg="#3498db", bg="#121212")
    web_lbl.place(relx=0.5, rely=0.65, anchor='center')

    # --- TYPING ANIMATION LOGIC ---
    def type_text(label, full_text, next_func=None, index=0):
        if index <= len(full_text):
            label.config(text=full_text[:index])
            splash.after(40, lambda: type_text(label, full_text, next_func, index + 1))
        elif next_func:
            splash.after(300, next_func)

    # --- APP LAUNCH SEQUENCE ---
    def launch_actual_app():
        splash.destroy()
        app_logic = QuotationApp(root)
        root.deiconify()
        root.attributes('-topmost', False) 
        root.state('zoomed') 
        root.lift()
        root.focus_force()

    # Animation Chain: Title -> ODM -> Slogan -> Powered By -> Web
    splash.after(500, lambda: type_text(title_lbl, "ODM Quick Books", 
      next_func=lambda: type_text(sub_lbl, "ODM-ONLINE", 
      next_func=lambda: type_text(slogan_lbl, "Precision in Every Proposal, Excellence in Every Quote.", 
      next_func=lambda: type_text(pow_lbl, "Powered by Orient Marketing", 
      next_func=lambda: type_text(web_lbl, "www.orientmarketing.com.pk"))))))

    # Total 15-18 seconds ka delay taake typing poori ho sake
    root.after(16000, launch_actual_app)
    
    root.mainloop()

if __name__ == "__main__":
    start_app()